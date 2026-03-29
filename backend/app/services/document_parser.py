import re
import hashlib
from typing import List, Optional, Iterable, Dict
from dataclasses import dataclass, field
from docx import Document as DocxDocument
import io
import datetime as dt

@dataclass
class DocumentMeta:
    document_id: str
    source_path: str
    title: str
    language: str = "vi"
    version: Optional[str] = None
    ingested_at: dt.datetime = field(default_factory=lambda: dt.datetime.now(dt.timezone.utc))

@dataclass
class Section:
    section_id: str
    heading_level: int
    title: str
    text: str
    parent_id: Optional[str]

@dataclass
class Chunk:
    chunk_id: str
    document_id: str
    section_id: str
    chunk_index: int
    text: str
    token_count: int
    metadata: dict

_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def _wtag(local: str) -> str:
    return f'{{{_W}}}{local}'


def _extract_paragraph_text(para, include_strikethrough: bool = True) -> str:
    """
    Extract text from a paragraph using raw OOXML lxml traversal.

    Handles three cases:
      - Normal runs (w:r / w:t)
      - Manual strikethrough runs (w:rPr contains w:strike or w:dstrike)
      - Word Track Changes deletions (w:del / w:delText) — invisible to python-docx API

    Struck-through text is wrapped in ~~...~~ sentinel markers.
    If include_strikethrough is False, such text is excluded entirely.
    """
    parts = []

    def _is_struck(rpr_elem):
        """Return True if a w:rPr element marks text as struck-through."""
        if rpr_elem is None:
            return False
        for tag_name in ('strike', 'dstrike'):
            el = rpr_elem.find(_wtag(tag_name))
            if el is not None:
                val = el.get(_wtag('val'), 'true')
                if val.lower() not in ('false', '0'):
                    return True
        return False

    for child in para._element:
        ctag = child.tag

        # Normal run: w:r
        if ctag == _wtag('r'):
            rpr = child.find(_wtag('rPr'))
            is_struck = _is_struck(rpr)
            text = ''.join(t.text or '' for t in child.findall(_wtag('t')))
            if not text:
                continue
            if is_struck:
                if include_strikethrough:
                    parts.append(f'~~{text}~~')
            else:
                parts.append(text)

        # Tracked deletion: w:del (Word Track Changes — completely skipped by python-docx)
        elif ctag == _wtag('del'):
            text = ''.join(dt.text or '' for dt in child.iter(_wtag('delText')))
            if not text:
                continue
            if include_strikethrough:
                parts.append(f'~~{text}~~')

        # Tracked insertion: w:ins — treat like normal text
        elif ctag == _wtag('ins'):
            for r in child.findall(_wtag('r')):
                rpr = r.find(_wtag('rPr'))
                is_struck = _is_struck(rpr)
                text = ''.join(t.text or '' for t in r.findall(_wtag('t')))
                if not text:
                    continue
                if is_struck:
                    if include_strikethrough:
                        parts.append(f'~~{text}~~')
                else:
                    parts.append(text)

    return ''.join(parts)


def _table_to_markdown(table) -> str:
    """Convert a python-docx Table to a Markdown table string."""
    rows = table.rows
    if not rows:
        return ""

    grid: List[List[str]] = []
    for row in rows:
        cells: List[str] = []
        prev_tc = None
        for cell in row.cells:
            if cell._tc is prev_tc:
                cells.append("")
            else:
                cell_text = " ".join(
                    p.text.strip() for p in cell.paragraphs if p.text.strip()
                )
                cells.append(cell_text.replace("|", "\\|"))
            prev_tc = cell._tc
        grid.append(cells)

    if not grid or all(not any(c for c in row) for row in grid):
        return ""

    max_cols = max(len(r) for r in grid)
    for r in grid:
        while len(r) < max_cols:
            r.append("")

    lines = [
        "| " + " | ".join(grid[0]) + " |",
        "| " + " | ".join("---" for _ in range(max_cols)) + " |",
    ]
    for row in grid[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


class DocLoader:
    """Loads DOCX or plain text files into normalized paragraphs."""
    # ... (existing methods load_from_bytes, _load_docx, _load_plain) ...
    def load_from_bytes(self, file_content: bytes, filename: str) -> List[str]:
        if filename.endswith(".docx"):
             return self._load_docx(file_content)
        elif filename.endswith(".pdf"):
             # Convert PDF to DOCX first using existing service
             try:
                 from app.services.document_service import document_service
                 docx_content = document_service.convert_pdf_to_docx(file_content)
                 return self._load_docx(docx_content)
             except Exception as e:
                 print(f"Error converting PDF via DocLoader: {e}")
                 # Fallback to plain text extraction (but broken structure)
                 # Or raise
                 raise e
        else:
             return self._load_plain(file_content)

    def _load_docx(self, file_content: bytes) -> List[str]:
        try:
            doc = DocxDocument(io.BytesIO(file_content))
        except Exception as exc:
            raise RuntimeError(f"Unable to parse DOCX: {exc}") from exc
        
        # Materialize body children list to keep lxml proxy references alive
        body_children = list(doc.element.body)
        element_pos = {id(child): idx for idx, child in enumerate(body_children)}

        items: list = []
        for para in doc.paragraphs:
            pos = element_pos.get(id(para._element))
            if pos is not None:
                items.append((pos, 'para', para))
        for table in doc.tables:
            pos = element_pos.get(id(table._tbl))
            if pos is not None:
                items.append((pos, 'table', table))

        items.sort(key=lambda x: x[0])

        paragraphs = []
        for _, item_type, item in items:
            if item_type == 'para':
                text = _extract_paragraph_text(item).strip()
                if text:
                    paragraphs.append(text)
            elif item_type == 'table':
                md = _table_to_markdown(item)
                if md:
                    paragraphs.append(md)
                    
        return paragraphs

    def _load_plain(self, file_content: bytes) -> List[str]:
        return [line.strip() for line in file_content.decode("utf-8", errors="ignore").splitlines() if line.strip()]

class SectionParser:
     # ... (existing implementation) ...
    HEADING_RE = re.compile(
        r"^(?:(?:Clause|Article|Section|Part|Chapter|Điều|Khoản|Mục)\s+)?(?:\((?P<paren_num>\d+)\)|(?P<num>\d+(?:\.\d+)*))\s*[:.]?\s+(?P<title>.+)",
        re.UNICODE | re.IGNORECASE,
    )

    def __init__(self):
        self._auto_counter = 1

    def parse(self, paragraphs: Iterable[str]) -> List[Section]:
        sections: List[Section] = []
        current_section: Optional[Section] = None
        self._auto_counter = 1

        for para in paragraphs:
            match = self.HEADING_RE.match(para)
            if match:
                section_id = match.group("num") or match.group("paren_num")
                level = section_id.count(".") + 1
                title = match.group("title").strip()
                parent_id = ".".join(section_id.split(".")[:-1]) or None
                current_section = Section(section_id, level, title, "", parent_id)
                sections.append(current_section)
            elif self._looks_like_heading(para):
                section_id = f"AUTO-{self._auto_counter}"
                self._auto_counter += 1
                current_section = Section(section_id, 1, para.rstrip(":. ").strip(), "", None)
                sections.append(current_section)
            elif current_section:
                current_section.text += ("\n" if current_section.text else "") + para
            else:
                fallback_id = "0"
                if not sections or sections[-1].section_id != fallback_id:
                    current_section = Section(fallback_id, 0, "Preface", para, None)
                    sections.append(current_section)
                else:
                    sections[-1].text += "\n" + para

        return [sec for sec in sections if sec.text.strip()]

    def _looks_like_heading(self, text: str) -> bool:
        clean = text.strip()
        if len(clean) < 5 or len(clean) > 140:
            return False
        words = clean.replace(":", "").replace(".", "").split()
        if not words or len(words) > 15:
            return False
        # ensure majority of words start uppercase or are all uppercase
        alpha_words = [w for w in words if any(ch.isalpha() for ch in w)]
        if not alpha_words:
            return False
        titlecase_ratio = sum(w[0].isupper() for w in alpha_words) / len(alpha_words)
        uppercase_ratio = sum(w.isupper() for w in alpha_words) / len(alpha_words)
        if uppercase_ratio > 0.6 or titlecase_ratio > 0.9:
            return True
        if clean.endswith(":"):
            return True
        return False

class Chunker:
    """Creates chunks per section with optional splitting by token length."""

    def __init__(self, max_tokens: int = 400):
        self.max_tokens = max_tokens

    def chunk(self, section: Section, document_id: str) -> List[Chunk]:
        text = section.text.strip()
        if not text:
            return []
        words = text.split()
        if len(words) <= self.max_tokens:
            return [self._create_chunk(document_id, section, 0, words)]

        chunks: List[Chunk] = []
        chunk_index = 0
        for start in range(0, len(words), self.max_tokens):
            window = words[start : start + self.max_tokens]
            if not window:
                continue
            chunks.append(self._create_chunk(document_id, section, chunk_index, window))
            chunk_index += 1
        return chunks

    def _create_chunk(
        self,
        document_id: str,
        section: Section,
        chunk_index: int,
        words: List[str],
    ) -> Chunk:
        text = " ".join(words).strip()
        chunk_id = hashlib.sha256(f"{document_id}-{section.section_id}-{chunk_index}-{text}".encode()).hexdigest()
        metadata = {
            "section_title": section.title,
            "heading_level": section.heading_level,
            "parent_section_id": section.parent_id,
        }
        return Chunk(
            chunk_id=chunk_id,
            document_id=document_id,
            section_id=section.section_id,
            chunk_index=chunk_index,
            text=text,
            token_count=len(words),
            metadata=metadata,
        )
