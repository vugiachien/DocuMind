from __future__ import annotations

"""
Document processing pipeline that loads text-only documents, parses numbered sections,
chunks them for embeddings, and writes vectors plus metadata to Milvus.
"""

import datetime as dt
import hashlib
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterable, List, Optional
import os
import logging
import tiktoken
from sentence_transformers import SentenceTransformer, util

logger = logging.getLogger(__name__)

from docx import Document as DocxDocument

from pymilvus import (
    Collection,
    CollectionSchema,
    DataType,
    FieldSchema,
    MilvusException,
    connections,
    utility,
)
from sentence_transformers import SentenceTransformer


@dataclass
class DocumentMeta:
    document_id: str
    source_path: str
    title: str
    language: str = "vi"
    version: Optional[str] = None
    ingested_at: dt.datetime = field(default_factory=lambda: dt.datetime.now(dt.timezone.utc))


@dataclass
class Section:
    section_id: str
    heading_level: int
    title: str
    text: str
    parent_id: Optional[str]
    comments: List[str] = field(default_factory=list)


@dataclass
class Chunk:
    chunk_id: str
    document_id: str
    section_id: str
    chunk_index: int
    text: str
    token_count: int
    metadata: dict


@dataclass
class Paragraph:
    text: str
    style: str
    is_bold: bool
    comments: List[str] = field(default_factory=list)


class DocLoader:
    """Loads DOCX or plain text files into normalized paragraphs with metadata."""

    def load(self, path: Path) -> List[Paragraph]:
        if not path.exists():
            raise FileNotFoundError(f"Document '{path}' does not exist")
        if not path.is_file():
            raise ValueError(f"Document '{path}' is not a file")
        suffix = path.suffix.lower()
        if suffix == ".docx":
            logger.info(f"Loading DOCX file: {path}")
            return self._load_docx(path)
        # PDF files should be converted to DOCX by Backend before reaching AI Service
        if suffix == ".pdf":
            raise ValueError(
                f"PDF files are not supported directly. "
                "Please convert to DOCX first using pdf2docx."
            )
        try:
            return self._load_plain(path)
        except OSError as exc:
            raise RuntimeError(f"Failed to read document '{path}': {exc}") from exc

    def _load_docx(self, path: Path) -> List[Paragraph]:
        try:
            doc = DocxDocument(path)
        except Exception as exc:
            raise RuntimeError(f"Unable to parse DOCX '{path}': {exc}") from exc
        
        # Parse comments xml
        comment_map = self._parse_comments(doc)
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # Materialize body children list to keep lxml proxy references alive
        # (prevents id() collisions from garbage-collected proxies)
        body_children = list(doc.element.body)
        element_pos = {id(child): idx for idx, child in enumerate(body_children)}

        # Collect paragraphs and tables with their positions
        items: list = []
        for para in doc.paragraphs:
            pos = element_pos.get(id(para._element))
            if pos is not None:
                items.append((pos, 'para', para))
        for table in doc.tables:
            pos = element_pos.get(id(table._tbl))
            if pos is not None:
                items.append((pos, 'table', table))

        items.sort(key=lambda x: x[0])

        paragraphs = []
        for _, item_type, item in items:
            if item_type == 'para':
                para = item
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"
                is_bold = False
                try:
                    if para.runs:
                        valid_runs = [r for r in para.runs if r.text.strip()]
                        if valid_runs:
                            bold_chars = sum(len(r.text) for r in valid_runs if r.bold)
                            total_chars = sum(len(r.text) for r in valid_runs)
                            if total_chars > 0 and (bold_chars / total_chars) > 0.8:
                                is_bold = True
                except Exception:
                    pass
                
                para_comments = []
                try:
                    refs = para._element.findall('.//w:commentReference', namespaces)
                    for ref in refs:
                        val = ref.get(f"{{{namespaces['w']}}}id")
                        if val and int(val) in comment_map:
                            para_comments.append(comment_map[int(val)])
                except Exception:
                    pass

                paragraphs.append(Paragraph(text=text, style=style_name, is_bold=is_bold, comments=para_comments))

            elif item_type == 'table':
                table_paragraphs = self._table_to_paragraphs(item)
                paragraphs.extend(table_paragraphs)
        
        logger.info(f"Loaded {len(paragraphs)} paragraphs (incl. tables) from {path}")
        return paragraphs

    # Maximum rows per table batch before splitting into multiple paragraphs
    TABLE_BATCH_SIZE = 15

    def _table_to_paragraphs(self, table) -> List[Paragraph]:
        """Convert a python-docx Table to one or more Paragraphs.
        
        - Small tables (≤ TABLE_BATCH_SIZE rows) → single Paragraph
        - Large tables → split into batches, each with the header row repeated
        - Wraps output in [TABLE DATA]...[/TABLE DATA] tags for LLM clarity
        """
        md = self._table_to_markdown(table)
        if not md:
            return []

        lines = md.split("\n")
        # lines[0] = header row, lines[1] = separator, lines[2:] = data rows
        if len(lines) <= 2:
            # Header-only table (no data rows)
            return [Paragraph(text=f"[TABLE DATA]\n{md}\n[/TABLE DATA]", style="Table", is_bold=False, comments=[])]

        header_line = lines[0]
        separator_line = lines[1]
        data_lines = lines[2:]

        if len(data_lines) <= self.TABLE_BATCH_SIZE:
            # Small table — single paragraph with context tag
            return [Paragraph(text=f"[TABLE DATA]\n{md}\n[/TABLE DATA]", style="Table", is_bold=False, comments=[])]

        # Large table — split into batches
        paragraphs = []
        for batch_idx in range(0, len(data_lines), self.TABLE_BATCH_SIZE):
            batch = data_lines[batch_idx:batch_idx + self.TABLE_BATCH_SIZE]
            batch_num = batch_idx // self.TABLE_BATCH_SIZE + 1
            total_batches = (len(data_lines) + self.TABLE_BATCH_SIZE - 1) // self.TABLE_BATCH_SIZE
            batch_md = "\n".join([header_line, separator_line] + batch)
            label = f"[TABLE DATA - Part {batch_num}/{total_batches}]"
            paragraphs.append(Paragraph(
                text=f"{label}\n{batch_md}\n[/TABLE DATA]",
                style="Table",
                is_bold=False,
                comments=[]
            ))
        
        logger.info(f"Split large table ({len(data_lines)} rows) into {len(paragraphs)} batches")
        return paragraphs

    @staticmethod
    def _table_to_markdown(table) -> str:
        """Convert a python-docx Table to a Markdown table string.
        
        Handles both horizontal merges (same _tc object) and vertical merges
        (w:vMerge XML attribute) to avoid duplicate content.
        """
        from docx.oxml.ns import qn
        
        rows = table.rows
        if not rows:
            return ""

        grid: List[List[str]] = []
        for row in rows:
            cells: List[str] = []
            prev_tc = None
            for cell in row.cells:
                tc = cell._tc
                if tc is prev_tc:
                    # Horizontally merged cell continuation — empty placeholder
                    cells.append("")
                else:
                    # Check for vertical merge continuation
                    vmerge = tc.find(qn('w:tcPr') + '/' + qn('w:vMerge'))
                    if vmerge is not None:
                        # vMerge with val="restart" = start of merge (show content)
                        # vMerge without val or val="continue" = continuation (skip)
                        val = vmerge.get(qn('w:val'))
                        if val != 'restart':
                            cells.append("")  # Continuation of vertical merge
                            prev_tc = tc
                            continue
                    
                    cell_text = " ".join(
                        p.text.strip() for p in cell.paragraphs if p.text.strip()
                    )
                    cells.append(cell_text.replace("|", "\\|"))
                prev_tc = tc
            grid.append(cells)

        if not grid or all(not any(c for c in row) for row in grid):
            return ""

        max_cols = max(len(r) for r in grid)
        for r in grid:
            while len(r) < max_cols:
                r.append("")

        lines = [
            "| " + " | ".join(grid[0]) + " |",
            "| " + " | ".join("---" for _ in range(max_cols)) + " |",
        ]
        for row in grid[1:]:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)

    def _parse_comments(self, doc) -> Dict[int, str]:
        """Extracts comment text mapped by ID from the docx package."""
        comments = {}
        try:
            # RELATIONSHIP_TYPE.COMMENTS is not always directly available easily, 
            # so we iterate rels to find the one pointing to comments.xml
            for rel in doc.part.rels.values():
                # Typical reltype for comments: .../relationships/comments
                if "comments" in rel.reltype and rel.target_part:
                    # Parse the XML of the comments part
                    from lxml import etree
                    xml_content = rel.target_part.blob
                    root = etree.fromstring(xml_content)
                    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    
                    # Find all w:comment elements
                    for comment_node in root.findall('.//w:comment', namespaces):
                        cid_str = comment_node.get(f"{{{namespaces['w']}}}id")
                        if cid_str:
                            # Extract all text nodes within the comment
                            text_nodes = comment_node.findall('.//w:t', namespaces)
                            full_text = "".join(node.text for node in text_nodes if node.text)
                            if full_text.strip():
                                comments[int(cid_str)] = full_text.strip()
        except Exception as e:
            # Fail silently on comment parsing, it's optional
            print(f"Warning: Failed to parse comments: {e}")
        return comments

    def _load_plain(self, path: Path) -> List[Paragraph]:
        lines = [line.strip() for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]
        return [Paragraph(text=t, style="Normal", is_bold=False, comments=[]) for t in lines]


class SectionParser:
    # Supported heading styles from Word
    HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle"}
    
    # Enhanced Regex for Vietnamese Legal Documents
    # Matches:
    # 1. (1) or 1. or 1.1 or 1.1.1
    # 2. Điều 1. or Điều 1:
    # 3. Chương I. or Chương I:
    # 4. Phần A. or Phần A:
    HEADING_RE = re.compile(
        r"^(?:"
        r"(?P<paren_num>\(\d+\))|"  # (1)
        r"(?P<num>\d+(?:\.\d+)*)|"   # 1. 1.1
        r"(?P<article>Điều\s+\d+)|"  # Điều 1
        r"(?P<chapter>Chương\s+[IVX0-9]+)|" # Chương I, Chương 1
        r"(?P<part>Phần\s+[A-Z])"    # Phần A
        r")\s*[:.]?\s+(?P<title>.+)?", # Title is optional if line ends with number
        re.UNICODE | re.IGNORECASE,
    )

    # Table-of-Contents header keywords (case-insensitive)
    TOC_HEADERS = {
        "TABLE OF CONTENTS", "TABLE OF CONTENT",
        "MỤC LỤC", "MỤC LỤC NỘI DUNG",
        "CONTENTS", "NỘI DUNG",
    }

    # TOC entry pattern: numbered/bulleted line ending with dots + page number
    # e.g. "1.   General Conditions..............5"
    # e.g. "ANNEX 1: PROJECT'S INFORMATION..........32"
    TOC_ENTRY_RE = re.compile(
        r'^.+\.{2,}\s*\d+\s*$',  # any text + 2+ dots + page number
        re.UNICODE,
    )

    def __init__(self):
        self._auto_counter = 1
        self._last_numbered_level = None

    def parse(self, paragraphs: Iterable[Paragraph]) -> List[Section]:
        sections: List[Section] = []
        current_section: Optional[Section] = None
        in_toc = False  # True while we are inside a TABLE OF CONTENTS block
        self._auto_counter = 1  # Reset per-parse so both docs get identical AUTO-N IDs
        self._last_numbered_level = None
        
        pending_heading_id = None
        pending_heading_level = None

        for para in paragraphs:
            text = para.text.strip()
            style = para.style

            # ── TABLE OF CONTENTS detection ──────────────────────────────────
            # Detect TOC header line
            if text.upper() in self.TOC_HEADERS:
                in_toc = True
                logger.debug(f"TOC header detected: '{text}' — entering skip mode")
                continue  # skip the header itself

            if in_toc:
                # Blank lines inside TOC block — stay in skip mode
                if not text:
                    continue
                # Exit TOC mode when we see a line that is NOT a TOC entry:
                # i.e. it does NOT end with dots+page-number AND it is non-empty.
                if self.TOC_ENTRY_RE.match(text):
                    logger.debug(f"  TOC entry skipped: '{text[:60]}'")
                    continue  # skip TOC entry
                else:
                    # First non-TOC line after the TOC block → resume normal parsing
                    in_toc = False
                    logger.debug(f"TOC block ended at: '{text[:60]}'")
                    # Fall through to normal processing below

            # ── End TOC detection ────────────────────────────────────────────

            # Priority 1: Check Word Styles
            is_style_heading = any(s in style for s in self.HEADING_STYLES)
            
            match = self.HEADING_RE.match(text)
            
            estimated_level = self._estimate_heading_level(match, current_section, is_style_heading)
            
            # Check if this paragraph is a standalone heading number (no title)
            if match and not match.group("title"):
                if match.group("article"):
                    pid = match.group("article").title()
                    plevel = 2
                elif match.group("chapter"):
                    pid = match.group("chapter")
                    if pid and pid[0].islower(): pid = pid[0].upper() + pid[1:]
                    plevel = 1
                elif match.group("part"):
                    pid = match.group("part").title()
                    plevel = 1
                else:
                    pid = match.group("num") or match.group("paren_num")
                    plevel = estimated_level if estimated_level else (pid.count(".") + 1 if pid else 1)
                
                if pending_heading_id:
                    logger.warning(f"Orphan heading detected: {pending_heading_id}")
                    # We no longer create empty sections. Just discard the orphaned id.
                    
                pending_heading_id = pid
                pending_heading_level = plevel
                continue
            
            if pending_heading_id:
                # If the NEXT paragraph is also a valid heading, we must NOT merge them!
                # We just drop the orphaned heading and process THIS paragraph as normal.
                if match or is_style_heading:
                    logger.warning(f"Orphan heading detected (followed by another heading): {pending_heading_id}")
                    pending_heading_id = None
                    # Fall through to process THIS heading normally
                else:
                    # Merge: this paragraph is the title/content for the pending heading
                    title = text
                    if len(title) > 100:
                        title = f"{pending_heading_id}"
                    current_section = Section(pending_heading_id, pending_heading_level, title, text, None, comments=para.comments)
                    sections.append(current_section)
                    self._last_numbered_level = pending_heading_level
                    pending_heading_id = None
                    continue
                
            if is_style_heading or match:
                # Determine ID and Title
                if match:
                    # Priority 2: Structure detected by Regex
                    if match.group("article"):
                        section_id = match.group("article").title() # Điều 1
                    elif match.group("chapter"):
                        # Don't use title() for Roman numerals (Chương II -> Chương Ii is wrong)
                        section_id = match.group("chapter") 
                        # Normalize first char just in case
                        if section_id and section_id[0].islower():
                            section_id = section_id[0].upper() + section_id[1:]
                    elif match.group("part"):
                         section_id = match.group("part").title()    # Phần A
                    else:
                        section_id = match.group("num") or match.group("paren_num")
                    
                    title = (match.group("title") or "").strip()
                    if not title:
                         title = text # If regex caught "Điều 1" but no title, use full text

                    # HEURISTIC: If title is too long, it's likely a list item text, not a heading
                    # e.g. "1. The Partner shall ensure that..." -> Content, not Heading
                    is_real_heading = True
                    if not match.group("article") and not match.group("chapter") and not match.group("part"):
                        # Only apply length check to numbered headings (1. or (1))
                        # Allow "Điều 1: Title" (Level 2) even if long? Usually titles are short.
                        if len(title) > 100:
                            is_real_heading = False
                        
                        # LOCAL SUB-NUMBERING: Numbers that look like local sub-items
                        # (like "1", "2", "2.1", "2.2") appearing under a style/heuristic
                        # heading (H-/AUTO-) are likely local sub-items, NOT top-level clauses.
                        # e.g. "Option 2" → "1. Down Payment" → "2. Payment based on..."
                        #       → "2.1. 1st Instalment" → "2.2. 2nd Instalment"
                        # Document-level numbers have large segments: "1.10", "1.11.1", "14.2"
                        num_id = match.group("num") or match.group("paren_num") or ""
                        if num_id and current_section:
                            segments = num_id.split(".")
                            # Local numbering: all segments ≤ 9 AND ≤ 2 levels deep
                            # e.g. "2" → [2], "2.1" → [2,1], "2.5" → [2,5] → local
                            # But "1.10" → [1,10] → segment 10 > 9 → document-level
                            # And "1.10.1" → 3 levels → document-level
                            all_small = all(s.isdigit() and int(s) <= 9 for s in segments if s)
                            shallow_depth = len(segments) <= 2
                            is_local_numbering = all_small and shallow_depth
                            is_parent_heuristic = (
                                current_section.section_id.startswith("H-") or
                                current_section.section_id.startswith("AUTO-")
                            )
                            if is_local_numbering and is_parent_heuristic:
                                is_real_heading = False
                            
                    if not is_real_heading:
                        # Treat as content
                        if current_section:
                            current_section.text += "\n" + para.text
                        else:
                             # Should not happen often if Preamble exists
                             pass 
                        # Continue to next paragraph
                        continue

                else:
                    # Style based only
                    section_id = f"H-{self._auto_counter}" 
                    self._auto_counter += 1
                    title = text
                
                # Determine level
                level = estimated_level
                if match and match.group("num"):
                    self._last_numbered_level = level
                
                # Initialize with the heading text itself so it's not empty/filtered out
                # and so the heading is included in embeddings
                # Comments on headings are rare but possible
                current_section = Section(section_id, level, title, text, None, comments=para.comments)
                sections.append(current_section)
                
            elif self._is_list_item(text, current_section):
                # Add to current section, don't create new section
                current_section.text += "\n" + text
                if para.comments:
                    current_section.comments.extend(para.comments)
                    
            elif self._looks_like_heading(para):
                # Only merge into current section if it has NO body content yet.
                # This handles: "Option 2" (empty) → "Down Payment..." (merges as content)
                # But does NOT merge across sections with real content, preventing
                # the entire document from collapsing into one section.
                if current_section and self._is_empty_section(current_section):
                    current_section.text += "\n" + text
                    if para.comments:
                        current_section.comments.extend(para.comments)
                else:
                    # Genuinely new heading section
                    section_id = f"AUTO-{self._auto_counter}"
                    self._auto_counter += 1
                    inferred_level = self._infer_auto_heading_level(current_section, sections)
                    current_section = Section(section_id, inferred_level, text.rstrip(":. ").strip(), text, None, comments=para.comments)
                    sections.append(current_section)
                
            elif current_section:
                current_section.text += ("\n" if current_section.text else "") + text
                if para.comments:
                    current_section.comments.extend(para.comments)
            else:
                fallback_id = "0"
                if not sections or sections[-1].section_id != fallback_id:
                    current_section = Section(fallback_id, 0, "Preface", text, None, comments=para.comments)
                    sections.append(current_section)
                else:
                    sections[-1].text += "\n" + text
                    if para.comments:
                        sections[-1].comments.extend(para.comments)

        if pending_heading_id:
            logger.warning(f"Orphan heading detected at end of doc: {pending_heading_id}")

        self._normalize_levels(sections)

        results = [sec for sec in sections if sec.text.strip()]
        logger.info(f"Parsed {len(results)} sections from {len(list(paragraphs)) if isinstance(paragraphs, list) else 'iterable'} paragraphs")
        return results

    def _is_empty_section(self, section: Section) -> bool:
        """Check if a section has no real body content beyond its title."""
        text_lines = [l.strip() for l in section.text.strip().split("\n") if l.strip()]
        # A section with 0 or 1 lines has no body content beyond its title
        return len(text_lines) <= 1

    def _is_list_item(self, text: str, current_section: Optional[Section]) -> bool:
        """Check if this paragraph is a list item belonging to current section."""
        list_pattern = re.match(r'^(\([ivx\d]+\)|\([a-z]\)|[a-z]\.)\s+', text, re.IGNORECASE)
        if list_pattern and current_section:
            section_text = current_section.text.strip()
            # Also check for introductory keywords
            intro_keywords = ["apply", "include", "following", "where", "if"]
            ends_with_colon = section_text.endswith(':')
            has_intro = any(kw in section_text.lower()[-50:] for kw in intro_keywords)
            
            if ends_with_colon or has_intro:
                return True
        return False

    def _looks_like_heading(self, para: Paragraph) -> bool:
        """Fallback heuristic for plain text"""
        text = para.text.strip()
        # Too short or too long is unlikely to be a valid heading for our purposes
        if len(text) < 5 or len(text) > 150:
            return False
            
        # Ignore common short preamble/filler words that are often capitalized or bold
        lower_text = text.lower().replace(":", "").strip()
        ignore_words = {"between", "and", "by and between", "no", "number", "date", "for", "hereby", "whereas"}
        if lower_text in ignore_words or lower_text.startswith("no:"):
            return False

        if para.is_bold: 
            return True # Bold lines usually headings if not matched by regex/style and pass length/ignore checks
            
        # Ends with colon (common in older docs)
        if text.endswith(':'):
            if len(text) > 60:
                return False
            continuation_patterns = [
                "shall not apply where",
                "shall apply where",
                "subject to",
                "provided that",
                "in the event that",
                "as follows",
                "the following",
                "include",
                "including"
            ]
            if any(p in text.lower() for p in continuation_patterns):
                return False
            return True

        words = text.replace(":", "").replace(".", "").split()
        if not words or len(words) > 30:
            return False
            
        # Heuristics: Capitalization
        if text.isupper(): return True
        
        # Title Case
        alpha_words = [w for w in words if any(ch.isalpha() for ch in w)]
        if not alpha_words:
            return False
        titlecase_ratio = sum(w[0].isupper() for w in alpha_words) / len(alpha_words)
        if titlecase_ratio > 0.85:
            return True
            
        return False

    def _estimate_heading_level(self, match: Optional[re.Match], current_section: Optional[Section], is_style_heading: bool) -> int:
        if match and match.group("num"):
            section_id = match.group("num")
            dot_count = section_id.count(".")
            if self._last_numbered_level is not None:
                return self._calculate_relative_level(section_id, dot_count)
            return max(1, dot_count)
        if match:
            if match.group("chapter"): return 1
            if match.group("part"): return 1
            if match.group("article"): return 2
            if match.group("paren_num"):
                if current_section: return current_section.heading_level + 1
                return 3
        if is_style_heading: return 1
        return 1

    def _calculate_relative_level(self, section_id: str, dot_count: int) -> int:
        return dot_count + 1

    def _infer_auto_heading_level(self, current_section: Optional[Section], sections: List[Section]) -> int:
        for i in range(len(sections) - 1, -1, -1):
            sec = sections[i]
            if sec.section_id.replace(".", "").isdigit():
                return sec.heading_level
        return 1

    def _normalize_levels(self, sections: List[Section]) -> None:
        if not sections: return
        numbered_sections = [(i, s) for i, s in enumerate(sections) if s.section_id.replace(".", "").isdigit()]
        if not numbered_sections: return
        
        min_dots = min(s.section_id.count(".") for _, s in numbered_sections)
        for idx, section in numbered_sections:
            dots = section.section_id.count(".")
            section.heading_level = dots - min_dots + 1
            
        for i, section in enumerate(sections):
            if section.section_id.startswith("AUTO-") or section.section_id.startswith("H-"):
                nearest_numbered_level = self._find_nearest_numbered_level(sections, i)
                section.heading_level = nearest_numbered_level
                
        self._populate_parent_ids(sections)
        
    def _find_nearest_numbered_level(self, sections: List[Section], current_idx: int) -> int:
        for i in range(current_idx - 1, -1, -1):
            if sections[i].section_id.replace(".", "").isdigit(): return sections[i].heading_level
        for i in range(current_idx + 1, len(sections)):
            next_id = sections[i].section_id
            if next_id.replace(".", "").isdigit(): 
                # If the next section is a sub-clause (e.g. 1.10.1), 
                # the current auto-heading is likely its parent.
                if next_id.endswith(".1"):
                    return max(1, sections[i].heading_level - 1)
                return sections[i].heading_level
        return 1

    def _populate_parent_ids(self, sections: List[Section]) -> None:
        level_stack: Dict[int, Section] = {}
        for section in sections:
            level = section.heading_level
            parent_level = level - 1
            if parent_level in level_stack:
                section.parent_id = level_stack[parent_level].section_id
            level_stack[level] = section
            levels_to_remove = [l for l in level_stack if l > level]
            for l in levels_to_remove: del level_stack[l]


class Chunker:
    """
    Sentence-aware chunker with Semantic and Adaptive Overlap enhancements.

    Splits section text into individual sentences first, then groups
    them into chunks that stay within *max_tokens* words.  Each chunk
    always ends at a sentence boundary, so 'Original Text' returned to
    the front-end is a readable, focused excerpt instead of a huge wall
    of text.

    Parameters
    ----------
    max_tokens : int
        Approximate maximum number of tokens per chunk.  Default 150.
    sentence_overlap : int
        Number of sentences from the end of the previous chunk to
        prepend to the next chunk for context continuity.  Default 1.
    semantic_threshold : float
        Min cosine similarity between current chunk and next sentence.
        If similarity is below this, the chunk is closed early. Default 0.3.
    """

    # Sentence-splitting regex:
    #   Split after  .  !  ?  when followed by whitespace + uppercase letter
    #   OR after  ;  when reasonably long (avoids splitting "e.g.; or")
    #   Does NOT split on abbreviations like "Mr.", "No.", or numeric "1."
    _SENT_RE = re.compile(
        r'(?<!\w\.\w.)'           # negative lookbehind: not abbrev like "U.S."
        r'(?<![A-Z][a-z]\.)'      # negative lookbehind: not "Dr."
        r'(?<!\d\.)'              # negative lookbehind: not "Article 1."
        r'(?<=[.!?])'             # split after sentence-ending punct
        r'\s+'                    # must be followed by whitespace
        r'(?=[A-ZÀ-Ỹa-z\("])',   # must be followed by letter / open paren
    )

    def __init__(self, max_tokens: int = 150, sentence_overlap: int = 1, semantic_threshold: float = 0.3):
        self.max_tokens = max_tokens
        self.sentence_overlap = sentence_overlap
        self.semantic_threshold = semantic_threshold
        # Use cl100k_base which is standard for current OpenAI models (ada-002, gpt-4)
        self.encoding = tiktoken.get_encoding("cl100k_base")

    def _count_tokens(self, text: str) -> int:
        """Count tokens using tiktoken instead of word split for better accuracy."""
        return len(self.encoding.encode(text))

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def chunk(self, section: Section, document_id: str, embedder: Optional['EmbeddingService'] = None) -> List[Chunk]:
        text = section.text.strip()
        if not text:
            return []

        sentences = self._split_into_sentences(text)
        if not sentences:
            return []

        # If the entire section fits in one chunk, keep it as-is
        total_tokens = sum(self._count_tokens(s) for s in sentences)
        if total_tokens <= self.max_tokens:
            return [self._make_chunk(document_id, section, 0, sentences)]

        return self._group_sentences(document_id, section, sentences, embedder)

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _split_into_sentences(self, text: str) -> List[str]:
        """
        Split *text* into a list of individual sentences.

        Strategy:
        1.  Detect markdown table rows (| ... |) and treat each row as a sentence.
        2.  Apply regex split on .!? boundaries for regular text.
        3.  Further split very long "sentences" on ; when they are very long.
        4.  Strip and drop empty results.
        """
        # Pre-process: if text contains table data, split table rows individually
        if '| --- |' in text or text.startswith('[TABLE DATA'):
            lines = text.split('\n')
            sentences: List[str] = []
            non_table_buffer: List[str] = []
            
            for line in lines:
                stripped = line.strip()
                if stripped.startswith('|') and stripped.endswith('|'):
                    # Flush any non-table text accumulated before this table row
                    if non_table_buffer:
                        buf_text = ' '.join(non_table_buffer)
                        sentences.extend(self._split_regular_text(buf_text))
                        non_table_buffer = []
                    # Skip separator rows (| --- | --- |)
                    if re.match(r'^\|[\s\-|]+\|$', stripped):
                        continue
                    sentences.append(stripped)
                elif stripped in ('[TABLE DATA]', '[/TABLE DATA]') or stripped.startswith('[TABLE DATA'):
                    # Context tags — include as-is for LLM awareness
                    sentences.append(stripped)
                else:
                    non_table_buffer.append(stripped)
            
            # Flush remaining non-table text
            if non_table_buffer:
                buf_text = ' '.join(non_table_buffer)
                sentences.extend(self._split_regular_text(buf_text))
            
            return sentences if sentences else [text.strip()]

        return self._split_regular_text(text)

    def _split_regular_text(self, text: str) -> List[str]:
        """Split regular (non-table) text into sentences using regex."""
        raw = self._SENT_RE.split(text)
        sentences: List[str] = []
        for s in raw:
            s = s.strip()
            if not s:
                continue
            # If a single "sentence" is still very long, try splitting on ;
            if self._count_tokens(s) > self.max_tokens:
                sub = re.split(r';\s+', s)
                for part in sub:
                    part = part.strip()
                    if part:
                        sentences.append(part)
            else:
                sentences.append(s)
        return sentences if sentences else [text.strip()]

    def _group_sentences(
        self,
        document_id: str,
        section: Section,
        sentences: List[str],
        embedder: Optional['EmbeddingService'] = None
    ) -> List[Chunk]:
        """Group sentences into token-limited chunks with sentence overlap and semantic cut."""
        
        # Pre-compute embeddings for all sentences in the section to batch process for semantic check
        sentence_embeddings = None
        if embedder:
            try:
                # Get embeddings for all sentences at once
                # self.model.encode returns a NumPy array
                embeddings_np = embedder.model.encode(sentences, show_progress_bar=False, convert_to_numpy=True)
                sentence_embeddings = embeddings_np
            except Exception as e:
                logger.warning(f"Failed to batch embed sentences for semantic check: {e}")
                
        chunks: List[Chunk] = []
        chunk_index = 0
        i = 0

        while i < len(sentences):
            window: List[str] = []
            token_count = 0
            
            # Start tracking the chunk semantic vector based on the first sentence
            current_chunk_embedding = sentence_embeddings[i] if (sentence_embeddings is not None and i < len(sentence_embeddings)) else None

            j = i
            while j < len(sentences):
                sent_tokens = self._count_tokens(sentences[j])
                
                # Check 1: Max Tokens limit
                if token_count + sent_tokens > self.max_tokens and window:
                    break  # chunk is full — stop before this sentence
                    
                # Check 2: Semantic Boundary (Only if we have a window > 0)
                if window and sentence_embeddings is not None and current_chunk_embedding is not None and j < len(sentence_embeddings):
                    next_sent_embedding = sentence_embeddings[j]
                    
                    # Compute cosine similarity
                    # util.cos_sim returns a matrix, we extract the scalar float
                    sim = util.cos_sim(current_chunk_embedding, next_sent_embedding).item()
                    
                    if sim < self.semantic_threshold:
                        logger.debug(f"Semantic cut triggered at sentence {j}: sim={sim:.3f} < {self.semantic_threshold:.3f}")
                        break

                # Update moving average of chunk embedding BEFORE appending,
                # so len(window) = n (previous items count), not n+1
                if sentence_embeddings is not None and current_chunk_embedding is not None and j < len(sentence_embeddings):
                    if len(window) > 0:  # Don't average if it's the first element
                        n = len(window)
                        current_chunk_embedding = (current_chunk_embedding * n + sentence_embeddings[j]) / (n + 1)
                    else:
                        current_chunk_embedding = sentence_embeddings[j]

                window.append(sentences[j])
                token_count += sent_tokens
                
                j += 1

            if not window:
                # Single sentence longer than max_tokens — include it anyway
                window = [sentences[i]]
                j = i + 1

            chunks.append(self._make_chunk(document_id, section, chunk_index, window))
            chunk_index += 1

            # Adaptive Overlap Logic
            # If the last string we added evaluates as "short" or "list header like", 
            # we should augment sentence_overlap
            dynamic_overlap = self.sentence_overlap
            last_sent = window[-1].strip()
            # If last sentence is less than 10 words, or ends with colon (introducing a list)
            if len(last_sent.split()) < 10 or last_sent.endswith(':'):
                 dynamic_overlap = max(1, self.sentence_overlap + 1)

            # Advance: move to j but step back by dynamic_overlap
            advance = max(1, j - i - dynamic_overlap)
            i += advance

        return chunks

    def _make_chunk(
        self,
        document_id: str,
        section: Section,
        chunk_index: int,
        sentences: List[str],
    ) -> Chunk:
        text = " ".join(sentences).strip()
        chunk_id = hashlib.sha256(
            f"{document_id}-{section.section_id}-{chunk_index}-{text}".encode()
        ).hexdigest()
        metadata = {
            "section_title": section.title,
            "heading_level": section.heading_level,
            "parent_section_id": section.parent_id,
        }
        return Chunk(
            chunk_id=chunk_id,
            document_id=document_id,
            section_id=section.section_id,
            chunk_index=chunk_index,
            text=text,
            token_count=self._count_tokens(text),
            metadata=metadata,
        )



class EmbeddingService:
    def __init__(self, model_name: str = "sentence-transformers/paraphrase-multilingual-mpnet-base-v2"):
        self.model = SentenceTransformer(model_name)

    def embed(self, texts: List[str]) -> List[List[float]]:
        embeddings = self.model.encode(texts, show_progress_bar=False, convert_to_numpy=True)
        return embeddings.tolist()

    @property
    def dimension(self) -> int:
        return self.model.get_sentence_embedding_dimension()


class DummyEntity:
    def __init__(self, data):
        self.data = data
    def get(self, key, default=None):
        return self.data.get(key, default)

class DummyHit:
    def __init__(self, data, distance):
        self.entity = DummyEntity(data)
        self.distance = distance

class ChromaVectorStore:
    def __init__(
        self,
        collection_name: str,
        embedding_dim: int,
        **kwargs
    ):
        self.collection_name = collection_name
        self.embedding_dim = embedding_dim
        
        import chromadb
        db_path = os.getenv("CHROMA_DB_PATH", os.path.join(os.getcwd(), "chroma_data"))
        os.makedirs(db_path, exist_ok=True)
        self.client = chromadb.PersistentClient(path=db_path)
        self.collection = self.client.get_or_create_collection(
            name=self.collection_name,
            metadata={"hnsw:space": "ip"}
        )
        print(f"ChromaVectorStore initialized at {db_path} for '{self.collection_name}'")

    def insert_chunks(self, chunks: List[Chunk], embeddings: List[List[float]]) -> None:
        if len(chunks) != len(embeddings):
            raise ValueError("Chunks and embeddings length mismatch")
            
        ids = [chunk.chunk_id for chunk in chunks]
        documents = [chunk.text[:8192] for chunk in chunks]
        metadatas = []
        for chunk in chunks:
            meta = {
                "document_id": chunk.document_id,
                "section_id": chunk.section_id,
                "chunk_index": chunk.chunk_index,
                "token_count": chunk.token_count,
                "metadata_json": json.dumps(chunk.metadata)
            }
            metadatas.append(meta)
            
        self.collection.add(
            ids=ids,
            embeddings=embeddings,
            metadatas=metadatas,
            documents=documents
        )

    def search(
        self,
        query: str,
        embedder: EmbeddingService,
        limit: int = 5,
        filter_expression: Optional[str] = None,
        output_fields: Optional[List[str]] = None,
    ):
        query_vector = embedder.embed([query])[0]
        
        where = None
        if filter_expression:
            import re
            match = re.search(r"document_id\s*==\s*['\"]([^'\"]+)['\"]", filter_expression)
            if match:
                where = {"document_id": match.group(1)}

        results = self.collection.query(
            query_embeddings=[query_vector],
            n_results=limit,
            where=where
        )
        
        ret = []
        if results and results["ids"] and len(results["ids"]) > 0:
            for i in range(len(results["ids"][0])):
                meta = results["metadatas"][0][i] if results["metadatas"] else {}
                doc = results["documents"][0][i] if results["documents"] else ""
                dist = results["distances"][0][i] if results["distances"] else 0.0
                
                data_dict = {
                    "document_id": meta.get("document_id", ""),
                    "section_id": meta.get("section_id", ""),
                    "chunk_index": meta.get("chunk_index", 0),
                    "metadata_json": meta.get("metadata_json", "{}"),
                    "chunk_text": doc,
                }
                ret.append(DummyHit(data_dict, dist))
        return [ret]

    def query_all_by_document_id(self, document_id: str) -> List[Dict]:
        results = self.collection.get(
            where={"document_id": document_id},
            include=["metadatas", "documents"]
        )
        
        ret = []
        if results and results["ids"]:
            for i in range(len(results["ids"])):
                meta = results["metadatas"][i]
                doc = results["documents"][i]
                ret.append({
                    "chunk_text": doc,
                    "section_id": meta.get("section_id", ""),
                    "chunk_index": meta.get("chunk_index", 0)
                })
                
        ret.sort(key=lambda x: (x.get("section_id", ""), x.get("chunk_index", 0)))
        return ret


class KnowledgeBaseBuilder:
    def __init__(
        self,
        embedder: EmbeddingService,
        store: ChromaVectorStore,
        loader: Optional[DocLoader] = None,
        parser: Optional[SectionParser] = None,
        chunker: Optional[Chunker] = None,
    ):
        self.loader = loader or DocLoader()
        self.parser = parser or SectionParser()
        self.chunker = chunker or Chunker()
        self.embedder = embedder
        self.store = store

    def ingest(self, path: Path, meta: DocumentMeta) -> int:
        paragraphs = self.loader.load(path)
        sections = self.parser.parse(paragraphs)

        all_chunks: List[Chunk] = []
        chunk_order = 0
        for order, section in enumerate(sections):
            # Pass embedder to trigger semantic chunking capability
            section_chunks = self.chunker.chunk(section, meta.document_id, self.embedder)
            for chunk in section_chunks:
                chunk.chunk_index = chunk_order
                chunk.metadata["section_order"] = order
                all_chunks.append(chunk)
                chunk_order += 1

        if not all_chunks:
            raise ValueError("No chunks generated from document")

        embeddings = self.embedder.embed([chunk.text for chunk in all_chunks])
        self.store.insert_chunks(all_chunks, embeddings)
        return len(all_chunks)