from typing import List, Tuple
from docx import Document
import io
import logging
import re

logger = logging.getLogger(__name__)


def _table_to_markdown(table) -> str:
    """Convert a python-docx Table to a Markdown table string."""
    rows = table.rows
    if not rows:
        return ""

    grid: List[List[str]] = []
    for row in rows:
        cells: List[str] = []
        prev_tc = None
        for cell in row.cells:
            if cell._tc is prev_tc:
                cells.append("")
            else:
                cell_text = " ".join(
                    p.text.strip() for p in cell.paragraphs if p.text.strip()
                )
                cells.append(cell_text.replace("|", "\\|"))
            prev_tc = cell._tc
        grid.append(cells)

    if not grid or all(not any(c for c in row) for row in grid):
        return ""

    max_cols = max(len(r) for r in grid)
    for r in grid:
        while len(r) < max_cols:
            r.append("")

    lines = [
        "| " + " | ".join(grid[0]) + " |",
        "| " + " | ".join("---" for _ in range(max_cols)) + " |",
    ]
    for row in grid[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


class DocumentService:
    def _normalize_text(self, text: str) -> str:
        """Normalize whitespace for text comparison"""
        import re
        # Convert all whitespace (including newlines and multiple spaces) to single space
        return re.sub(r'\s+', ' ', text).strip()

    def _find_similar_texts(self, doc, target_text: str, threshold: int = 70) -> List[str]:
        """Find texts in document that are similar to target_text"""
        from fuzzywuzzy import fuzz
        
        normalized_target = self._normalize_text(target_text)
        candidates = []
        
        # Collect all paragraphs and cell texts
        all_texts = []
        for p in doc.paragraphs:
            if p.text.strip():
                all_texts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            all_texts.append(p.text)
                            
        # Score them
        for text in all_texts:
            normalized_text = self._normalize_text(text)
            # Use partial ratio to find if target is IN the text
            if len(normalized_target) < len(normalized_text):
                ratio = fuzz.partial_ratio(normalized_target, normalized_text)
            else:
                ratio = fuzz.ratio(normalized_target, normalized_text)
                
            if ratio >= threshold:
                candidates.append((text, ratio))
                
        # Sort by score desc
        candidates.sort(key=lambda x: x[1], reverse=True)
        return [c[0] for c in candidates[:3]]

    def replace_text_in_docx(
        self, 
        file_content: bytes, 
        target_text: str, 
        replacement_text: str, 
        fuzzy_match: bool = True
    ) -> Tuple[bytes, dict]:
        """
        Replace target_text with replacement_text in a DOCX file content.
        Supports exact match, normalized match, and fuzzy match.
        
        Returns:
            (file_content_bytes, metadata_dict)
            
        Raises:
            TextReplacementError: If text cannot be found
        """
        from app.services.exceptions import TextReplacementError
        from fuzzywuzzy import fuzz
        
        try:
            logger.info(f"🔍 [DOCX Replace] Starting replacement...")
            logger.info(f"   Target: {target_text[:50]}...")
            
            # Load document
            doc = Document(io.BytesIO(file_content))
            
            replaced = False
            metadata = {
                "success": False,
                "method": None,
                "confidence": 0.0,
                "formatting_preserved": False, # Placeholder for Phase 2
                "warnings": []
            }
            
            # Helper to process paragraph
            def process_paragraph(paragraph):
                nonlocal replaced
                
                # Strategy 1: Exact Match
                if target_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(target_text, replacement_text)
                    logger.info(f"✅ Exact match found")
                    metadata["method"] = "exact"
                    metadata["confidence"] = 1.0
                    metadata["success"] = True
                    return True

                # Strategy 2: Normalized Match
                normalized_target = self._normalize_text(target_text)
                normalized_para = self._normalize_text(paragraph.text)
                
                if normalized_target in normalized_para:
                    # Found roughly, but need to replace in original text
                    # This is naive replacement (replaces whole paragraph content if it matches closely)
                    # For now, if we match via normalization, we might struggle to replace EXACT substring 
                    # while keeping other parts if they differ only by whitespace.
                    # Simplest approach for Phase 1: Try to replace strict if possible, 
                    # else if paragraph text IS the target (roughly), replace whole text.
                    
                    if normalized_target == normalized_para:
                        paragraph.text = replacement_text
                        logger.info(f"✅ Normalized full match found")
                        metadata["method"] = "normalized"
                        metadata["confidence"] = 0.95
                        metadata["success"] = True
                        return True
                        
                # Strategy 3: Fuzzy Match (if enabled)
                if fuzzy_match:
                    score = fuzz.ratio(normalized_target, normalized_para)
                    if score >= 85:
                        logger.info(f"✅ Fuzzy match found (score: {score})")
                        # Warning: This replaces WHOLE paragraph if it's a close match
                        # Ideally needed more granular replacement
                        paragraph.text = replacement_text
                        metadata["method"] = "fuzzy"
                        metadata["confidence"] = score / 100.0
                        metadata["success"] = True
                        metadata["warnings"].append(f"Used fuzzy match (score: {score})")
                        return True
                        
                return False

            # 1. Search in paragraphs
            for paragraph in doc.paragraphs:
                if process_paragraph(paragraph):
                    replaced = True
                    break # Stop after first replacement for now (finding usually targets specific instance)
            
            # 2. Search in tables if not found
            if not replaced:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if process_paragraph(paragraph):
                                    replaced = True
                                    break
                        if replaced: break
                    if replaced: break
            
            if not replaced:
                logger.warning(f"❌ Text NOT FOUND")
                # Find suggestions
                suggestions = self._find_similar_texts(doc, target_text)
                raise TextReplacementError(
                    reason="Target text not found in document",
                    target_text=target_text,
                    suggestions=suggestions
                )
                
            # Save
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output.read(), metadata
            
        except TextReplacementError:
            raise
        except Exception as e:
            logger.error(f"❌ [DOCX Replace] ERROR: {e}")
            raise Exception(f"Document processing error: {str(e)}")

    def _replace_text_in_paragraph_preserve_format(self, paragraph, target_text: str, replacement_text: str) -> bool:
        """
        Replace target_text with replacement_text inside a paragraph while preserving
        the formatting of the first run that contains the match.
        
        Works by:
        1. Concatenating all run texts to find the match position
        2. Mapping character positions back to runs
        3. Replacing text in runs, keeping the first matched run's formatting
        """
        # Build full text from runs and map char positions to (run_index, offset_in_run)
        runs = paragraph.runs
        if not runs:
            return False
            
        full_text = ''.join(r.text for r in runs)
        match_start = full_text.find(target_text)
        if match_start == -1:
            return False
        
        match_end = match_start + len(target_text)
        
        # Map each character position to (run_index, offset_in_run)
        char_to_run = []
        for ri, run in enumerate(runs):
            for ci in range(len(run.text)):
                char_to_run.append((ri, ci))
        
        # Find which runs are involved
        start_run_idx, start_offset = char_to_run[match_start]
        end_run_idx, end_offset = char_to_run[match_end - 1]
        
        # Put the replacement text in the first matched run (preserves its formatting)
        first_run = runs[start_run_idx]
        prefix = first_run.text[:start_offset]
        
        last_run = runs[end_run_idx]
        suffix = last_run.text[end_offset + 1:]
        
        if start_run_idx == end_run_idx:
            # Match is within a single run
            first_run.text = prefix + replacement_text + suffix
        else:
            # Match spans multiple runs — put replacement in first, clear the rest
            first_run.text = prefix + replacement_text
            for ri in range(start_run_idx + 1, end_run_idx):
                runs[ri].text = ''
            last_run.text = suffix
        
        return True

    def replace_multiple_texts_in_docx(self, file_content: bytes, replacements: List[Tuple[str, str]]) -> Tuple[bytes, List[str]]:
        """
        Replace multiple target texts with replacements in a DOCX file in ONE pass.
        Preserves formatting by operating at the run level.
        Supports exact match, normalized match, and fuzzy match.
        
        Args:
            file_content: DOCX file as bytes
            replacements: List of tuples [(target1, replacement1), (target2, replacement2), ...]
        
        Returns:
            Tuple[bytes, List[str]]: (Modified DOCX file as bytes, List of target texts that failed to match)
        """
        from fuzzywuzzy import fuzz
        
        try:
            # Load document ONCE
            doc = Document(io.BytesIO(file_content))
            
            replaced_count = 0
            failed_targets = []
            
            for target_text, replacement_text in replacements:
                normalized_target = self._normalize_text(target_text)
                replaced_this = False
                
                # Helper function to try replacement on a paragraph
                def try_replace_paragraph(paragraph):
                    nonlocal replaced_this, replaced_count
                    
                    # Strategy 1: Exact match — format-preserving run-level replacement
                    if target_text in paragraph.text:
                        if self._replace_text_in_paragraph_preserve_format(paragraph, target_text, replacement_text):
                            logger.info(f"✅ Exact match for: {target_text[:30]}...")
                            replaced_this = True
                            replaced_count += 1
                            return True
                    
                    # Strategy 2: Normalized match — only when target ≈ entire paragraph
                    normalized_para = self._normalize_text(paragraph.text)
                    if normalized_target in normalized_para:
                        score = fuzz.ratio(normalized_target, normalized_para)
                        if score >= 85:
                            # Target is essentially the whole paragraph — safe to replace all runs
                            if paragraph.runs:
                                paragraph.runs[0].text = replacement_text
                                for r in paragraph.runs[1:]:
                                    r.text = ''
                            else:
                                paragraph.text = replacement_text
                            logger.info(f"✅ Normalized match (score: {score}) for: {target_text[:30]}...")
                            replaced_this = True
                            replaced_count += 1
                            return True
                    
                    # Strategy 3: Fuzzy match — only when target ≈ entire paragraph (strict threshold)
                    score = fuzz.ratio(normalized_target, normalized_para)
                    if score >= 88:
                        if paragraph.runs:
                            paragraph.runs[0].text = replacement_text
                            for r in paragraph.runs[1:]:
                                r.text = ''
                        else:
                            paragraph.text = replacement_text
                        logger.info(f"✅ Fuzzy match (score: {score}) for: {target_text[:30]}...")
                        replaced_this = True
                        replaced_count += 1
                        return True
                    
                    return False
                
                # Search in paragraphs
                for paragraph in doc.paragraphs:
                    if try_replace_paragraph(paragraph) and replaced_this:
                        break  # Found and replaced, move to next target
                
                # Search in tables if not found
                if not replaced_this:
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    if try_replace_paragraph(paragraph) and replaced_this:
                                        break
                                if replaced_this:
                                    break
                            if replaced_this:
                                break
                        if replaced_this:
                            break
                
                if not replaced_this:
                    logger.warning(f"⚠️ Target text NOT FOUND: {target_text[:50]}...")
                    failed_targets.append(target_text)
            
            logger.info(f"📊 Batch replacement summary: {replaced_count}/{len(replacements)} texts replaced. Failed: {len(failed_targets)}")
            
            # Save ONCE
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output.read(), failed_targets
            
        except Exception as e:
            logger.error(f"Error in batch DOCX manipulation: {e}")
            raise e

    def convert_pdf_to_docx(self, pdf_content: bytes) -> bytes:
        """
        Convert PDF file to DOCX format using pdf2docx library.
        
        Args:
            pdf_content: PDF file as bytes
        
        Returns:
            DOCX file as bytes
            
        Raises:
            Exception: If conversion fails
        """
        try:
            import io
            import tempfile
            import os
            from pdf2docx import Converter
            
            # pdf2docx requires file paths, so use temp files
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_temp:
                pdf_temp.write(pdf_content)
                pdf_temp_path = pdf_temp.name
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_temp:
                docx_temp_path = docx_temp.name
            
            try:
                # Convert PDF to DOCX
                cv = Converter(pdf_temp_path)
                cv.convert(docx_temp_path, start=0, end=None)
                cv.close()
                
                # Read converted DOCX
                with open(docx_temp_path, 'rb') as f:
                    docx_content = f.read()
                
                return docx_content
                
            finally:
                # Cleanup temp files
                if os.path.exists(pdf_temp_path):
                    os.unlink(pdf_temp_path)
                if os.path.exists(docx_temp_path):
                    os.unlink(docx_temp_path)
                    
        except Exception as e:
            logger.error(f"Error converting PDF to DOCX: {e}")
            raise e

    def extract_text_pdf(self, file_content: bytes) -> str:
        """
        Extract full text from PDF bytes using pdfplumber.
        """
        import pdfplumber
        try:
            full_text = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
                        
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise e

    def extract_text(self, file_content: bytes) -> str:
        """
        Extract full text from file bytes (DOCX or PDF).
        Strikethrough runs in DOCX are preserved as ~~text~~ sentinel markers.
        Tables are converted to Markdown format to preserve spatial relationships.
        """
        # Check for PDF magic bytes (%PDF)
        if file_content.startswith(b'%PDF'):
            return self.extract_text_pdf(file_content)

        try:
            from app.services.document_parser import _extract_paragraph_text

            doc = Document(io.BytesIO(file_content))
            full_text = []

            # Materialize body children list to keep lxml proxy references alive
            body_children = list(doc.element.body)
            element_pos = {id(child): idx for idx, child in enumerate(body_children)}

            items: list = []
            for para in doc.paragraphs:
                pos = element_pos.get(id(para._element))
                if pos is not None:
                    items.append((pos, 'para', para))
            for table in doc.tables:
                pos = element_pos.get(id(table._tbl))
                if pos is not None:
                    items.append((pos, 'table', table))

            items.sort(key=lambda x: x[0])

            for _, item_type, item in items:
                if item_type == 'para':
                    full_text.append(_extract_paragraph_text(item))
                elif item_type == 'table':
                    md = _table_to_markdown(item)
                    if md:
                        full_text.append(md)

            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise e

    def compare_texts(self, text1: str, text2: str) -> dict:
        """
        Compare two text strings using Standard Paragraph-Based Diff.
        - Line-by-line diff for paragraph detection (removed / added / unchanged)
        - When removed immediately followed by added → word-level diff for inline highlights
        - Strikethrough sentinels ~~...~~ are expanded into 'strikethrough' chunks
        """
        import difflib
        import re

        STRIKE_RE = re.compile(r'(~~.+?~~)')

        def split_strike_chunks(text: str, base_type: str) -> list:
            parts = STRIKE_RE.split(text)
            out = []
            for part in parts:
                if not part:
                    continue
                if part.startswith('~~') and part.endswith('~~'):
                    out.append({'type': 'strikethrough', 'content': part[2:-2]})
                else:
                    out.append({'type': base_type, 'content': part})
            return out

        def compute_inline_diff(old_text: str, new_text: str):
            """
            Word-token diff between two paragraph strings.
            Returns a single interleaved list of parts:
              - unchanged_part: text identical in both (black)
              - removed_part: text only in old (red strikethrough)
              - added_part: text only in new (green)
            """
            tokens_old = re.split(r'(\s+)', old_text)
            tokens_new = re.split(r'(\s+)', new_text)

            sm = difflib.SequenceMatcher(None, tokens_old, tokens_new, autojunk=False)
            parts: list = []

            for op, i1, i2, j1, j2 in sm.get_opcodes():
                old_tok = ''.join(tokens_old[i1:i2])
                new_tok = ''.join(tokens_new[j1:j2])

                if op == 'equal':
                    if old_tok:
                        parts.append({'type': 'unchanged_part', 'content': old_tok})
                else:  # replace / delete / insert
                    if old_tok:
                        parts.append({'type': 'removed_part', 'content': old_tok})
                    if new_tok:
                        parts.append({'type': 'added_part', 'content': new_tok})

            return parts

        # ── Step 1: Line-level ndiff ─────────────────────────────────────────
        lines1 = text1.splitlines(keepends=True)
        lines2 = text2.splitlines(keepends=True)
        raw_diff = list(difflib.ndiff(lines1, lines2))

        # ── Step 2: Group consecutive same-type lines into chunks ─────────────
        raw_chunks: list = []
        current: dict = {'type': None, 'content': ''}

        for line in raw_diff:
            if line.startswith('?'):
                continue
            code = line[0]
            content = line[2:]

            if code == ' ':
                t = 'unchanged'
            elif code == '-':
                t = 'removed'
            elif code == '+':
                t = 'added'
            else:
                continue

            if current['type'] == t:
                current['content'] += content
            else:
                if current['type']:
                    raw_chunks.append(dict(current))
                current = {'type': t, 'content': content}

        if current['type']:
            raw_chunks.append(current)

        # ── Step 3: Pair removed+added chunks → word-level inline diff ────────
        enhanced: list = []
        i = 0
        while i < len(raw_chunks):
            chunk = raw_chunks[i]
            if (chunk['type'] == 'removed'
                    and i + 1 < len(raw_chunks)
                    and raw_chunks[i + 1]['type'] == 'added'):
                flat_parts = compute_inline_diff(
                    chunk['content'], raw_chunks[i + 1]['content']
                )
                enhanced.append({'type': 'replaced', 'content': chunk['content'] + raw_chunks[i + 1]['content'], 'parts': flat_parts})
                i += 2
            else:
                enhanced.append(chunk)
                i += 1

        # ── Step 4: Expand ~~strikethrough~~ sentinels ────────────────────────
        expanded: list = []
        for chunk in enhanced:
            content = chunk.get('content', '')
            if 'parts' in chunk:
                # Chunk has word-level parts — clean ~~ sentinels within each part
                # rather than discarding parts[] and breaking layout with split_strike_chunks
                new_parts = []
                for part in chunk['parts']:
                    cleaned = STRIKE_RE.sub(lambda m: m.group(0)[2:-2], part['content'])
                    new_parts.append({'type': part['type'], 'content': cleaned})
                expanded.append({'type': chunk['type'], 'content': content, 'parts': new_parts})
            elif '~~' in content:
                # Plain chunk (no word-level parts) — expand into sub-chunks
                expanded.extend(split_strike_chunks(content, chunk['type']))
            else:
                expanded.append({'type': chunk['type'], 'content': content})

        return {'diff': expanded}


    def compare_documents(self, old_content: bytes, new_content: bytes) -> dict:
        """
        Compare two DOCX files and return diff chunks.
        """
        text1 = self.extract_text(old_content)
        text2 = self.extract_text(new_content)
        
        return self.compare_texts(text1, text2)

    def extract_docx_comments(self, file_content: bytes) -> list:
        """
        Extract comments from DOCX file along with the text they reference.
        
        Args:
            file_content: DOCX file as bytes
            
        Returns:
            List of comments with structure:
            [
                {
                    'id': str,
                    'author': str,
                    'date': str,
                    'text': str,
                    'quote': str  # The text in the document that this comment refers to
                },
                ...
            ]
        """
        try:
            from zipfile import ZipFile
            import xml.etree.ElementTree as ET
            
            comments_map = {}
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            with ZipFile(io.BytesIO(file_content)) as docx_zip:
                # 1. Parse comments.xml to get metadata and content
                if 'word/comments.xml' not in docx_zip.namelist():
                    return []
                
                comments_xml = docx_zip.read('word/comments.xml')
                root_comments = ET.fromstring(comments_xml)
                
                for comment in root_comments.findall('.//w:comment', ns):
                    cid = comment.get(f'{{{ns["w"]}}}id')
                    
                    # Extract comment text
                    text_parts = []
                    for text_elem in comment.findall('.//w:t', ns):
                        if text_elem.text:
                            text_parts.append(text_elem.text)
                    
                    comments_map[cid] = {
                        'id': cid,
                        'author': comment.get(f'{{{ns["w"]}}}author') or 'Unknown',
                        'date': comment.get(f'{{{ns["w"]}}}date') or '',
                        'text': ''.join(text_parts),
                        'quote': '' # Will be filled from document.xml
                    }
                
                # 2. Parse document.xml to extract referenced text
                if 'word/document.xml' in docx_zip.namelist():
                    doc_xml = docx_zip.read('word/document.xml')
                    root_doc = ET.fromstring(doc_xml)
                    
                    # Track active ranges: cid -> list of text parts
                    active_ranges = {}
                    
                    # Recursive function to traverse XML and extract text
                    def traverse(node):
                        tag = node.tag
                        
                        # Handle Comment Range Start
                        if tag == f'{{{ns["w"]}}}commentRangeStart':
                            cid = node.get(f'{{{ns["w"]}}}id')
                            if cid in comments_map:
                                active_ranges[cid] = []
                        
                        # Handle Comment Range End
                        elif tag == f'{{{ns["w"]}}}commentRangeEnd':
                            cid = node.get(f'{{{ns["w"]}}}id')
                            if cid in active_ranges:
                                # Join collected text
                                raw_quote = ''.join(active_ranges[cid])
                                # Normalize whitespace to match mammoth HTML rendering
                                # (collapse multiple spaces, newlines to single space)
                                import re
                                normalized_quote = re.sub(r'\s+', ' ', raw_quote).strip()
                                comments_map[cid]['quote'] = normalized_quote
                                del active_ranges[cid]
                        
                        # Handle Text Nodes
                        elif tag == f'{{{ns["w"]}}}t':
                            text = node.text or ""
                            # Append text to ALL currently active ranges
                            for cid in active_ranges:
                                active_ranges[cid].append(text)
                                
                        # Recurse children
                        for child in node:
                            traverse(child)
                            
                    traverse(root_doc)

            # Return list sorted by ID (usually order of creation, but improved: could sort by position if needed)
            # For now, simplistic list return
            return list(comments_map.values())
            
        except Exception as e:
            logger.error(f"Error extracting comments from DOCX: {e}")
            return []

    def convert_docx_to_html(self, docx_content: bytes) -> str:
        """
        Convert DOCX to HTML using Python mammoth library.
        Server-side rendering for instant frontend previews.
        
        Args:
            docx_content: DOCX file as bytes
            
        Returns:
            HTML string with embedded styles
        
        ✅ OPTIMIZATION: Eliminates client-side mammoth.js bottleneck (2-7s → < 100ms)
        """
        try:
            import mammoth
            import io
            
            logger.info(f"🔄 [HTML Conversion] Converting DOCX to HTML...")
            
            # Convert bytes to file-like object
            docx_file = io.BytesIO(docx_content)
            
            # Convert with mammoth
            result = mammoth.convert_to_html(docx_file)
            html = result.value
            
            # Wrap in styled container (match frontend styling)
            styled_html = f"""
<div style="font-family: 'Times New Roman', serif; font-size: 16px; line-height: 1.6; padding: 20px; color: #262626;">
    {html}
</div>
"""
            
            logger.info(f"✅ [HTML Conversion] Success! Generated {len(styled_html):,} chars of HTML")
            
            return styled_html
            
        except ImportError:
            logger.error(f"❌ [HTML Conversion] ERROR: mammoth library not installed")
            logger.error(f"   Run: pip install mammoth")
            raise Exception("mammoth library required for HTML conversion")
        except Exception as e:
            logger.error(f"❌ [HTML Conversion] ERROR: {e}")
            raise

    def convert_docx_to_html_with_comments(self, docx_content: bytes) -> str:
        """
        Convert DOCX to HTML using custom XML-based converter that preserves
        comment range markers as <span data-comment-id="X"> elements.

        Falls back to mammoth if the custom converter fails.
        """
        try:
            from app.services.docx_html_converter import DocxHtmlConverter

            logger.info("🔄 [HTML+Comments] Converting DOCX with custom converter...")
            html = DocxHtmlConverter.convert(docx_content)
            logger.info(f"✅ [HTML+Comments] Success! Generated {len(html):,} chars")
            return html

        except Exception as e:
            logger.warning(f"⚠️ [HTML+Comments] Custom converter failed ({e}), falling back to mammoth")
            return self.convert_docx_to_html(docx_content)

document_service = DocumentService()


