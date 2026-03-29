from typing import List, Dict, Any, Optional
from docx import Document
from docx.text.paragraph import Paragraph
import re
import io
import logging

logger = logging.getLogger(__name__)

class ContractModifier:
    def __init__(self):
        # Ported from ai_service/document_pipeline.py SectionParser
        self.HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle"}
        self.HEADING_RE = re.compile(
            r"^(?:"
            r"(?P<paren_num>\(\d+\))|"  # (1)
            r"(?P<num>\d+(?:\.\d+)*)|"   # 1. 1.1
            r"(?P<article>Điều\s+\d+)|"  # Điều 1
            r"(?P<chapter>Chương\s+[IVX0-9]+)|" # Chương I, Chương 1
            r"(?P<part>Phần\s+[A-Z])"    # Phần A
            r")\s*[:.]?\s+(?P<title>.+)?", 
            re.UNICODE | re.IGNORECASE,
        )

    def process_modifications(self, file_content: bytes, actions: List[Dict[str, Any]]) -> bytes:
        """
        Main entry point: Load -> Mutate -> Reindex -> Save
        """
        doc = Document(io.BytesIO(file_content))
        
        # 1. Map Structure
        # We need a stable reference to sections to find targets/anchors
        # Map: section_id (or simple sequential index if ID missing) -> Paragraph
        # NOTE: The current AI returns "sec_01", but we need to know what that maps to.
        # Implication: We need to re-scan the doc to match the IDs the AI *saw*.
        # Strategy: The AI saw "Title". We match by Title or sequence.
        # Better: Match by "Header Text" fuzzy match if IDs aren't persistent.
        # Current System: `ContractAnalysis` assigns IDs based on order (sec_0, sec_1).
        # We must replicate that ordering to be sure.
        
        sections = self._map_sections(doc)
        
        # 2. Apply Actions
        # Sort actions: MODIFY first, then INSERT (to preserve anchors)
        # However, if we insert *before* an anchor, references might shift?
        # Safe strategy:
        # - Process MODIFY actions first (in place updates).
        # - Process INSERT actions next (insertions).
        
        # Filter actions
        modifications = [a for a in actions if a.get('action') == 'MODIFY']
        insertions = [a for a in actions if a.get('action') == 'INSERT']
        
        for mod in modifications:
            self._apply_modify(doc, sections, mod)
            
        for ins in insertions:
            self._apply_insert(doc, sections, ins)

        # 3. Re-index
        self._reindex_document(doc)
        
        # 4. Save
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _map_sections(self, doc) -> List[Dict]:
        """
        Scan doc to find Headings/Sections that match the AI's view.
        Returns ordered list of { 'id': str, 'paragraph': Paragraph, 'text': str }
        """
        mapped = []
        _auto_counter = 1
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
                
            style = p.style.name if p.style else ""
            is_style_heading = any(s in style for s in self.HEADING_STYLES)
            match = self.HEADING_RE.match(text)
            
            section_id = None
            
            if is_style_heading or match:
                if match:
                    if match.group("article"):
                        section_id = match.group("article").title() # Điều 1
                    elif match.group("chapter"):
                        section_id = match.group("chapter")
                        if section_id and section_id[0].islower():
                            section_id = section_id[0].upper() + section_id[1:]
                    elif match.group("part"):
                         section_id = match.group("part").title()
                    else:
                        section_id = match.group("num") or match.group("paren_num")
                else:
                    section_id = f"H-{_auto_counter}"
                    _auto_counter += 1
            elif self._looks_like_heading(p):
                 section_id = f"AUTO-{_auto_counter}"
                 _auto_counter += 1
                 
            if section_id:
                 mapped.append({
                     "id": section_id,
                     "paragraph": p,
                     "text": text
                 })
                 
        return mapped

    def _looks_like_heading(self, para) -> bool:
        """Fallback heuristic for plain text (ported from SectionParser)"""
        # Check bold via runs
        is_bold = False
        if para.runs and para.runs[0].bold:
            is_bold = True
            
        if is_bold: 
            return True 
            
        text = para.text.strip()
        if len(text) < 5 or len(text) > 140:
            return False
        words = text.replace(":", "").replace(".", "").split()
        if not words or len(words) > 15:
            return False
            
        if text.isupper(): return True
        
        alpha_words = [w for w in words if any(ch.isalpha() for ch in w)]
        if not alpha_words:
            return False
        titlecase_ratio = sum(w[0].isupper() for w in alpha_words) / len(alpha_words)
        if titlecase_ratio > 0.9:
            return True
            
        return False

    def _apply_modify(self, doc, sections: List[Dict], action: Dict):
        """
        action: { target_id: "sec_XX", new_content: "..." }
        """
        target_id = action.get('target_id')
        
        # Find target
        # Problem: AI might return "sec_05" meaning "Điều 5".
        # If we rely on _map_sections index, we must trust the AI counted correctly.
        # Fallback: Match by "Original Text" (if provided) is safer?
        # User Plan says: "Find by ID".
        # Let's try to find by ID first (assuming sequential integrity).
        
        # Find target with fuzzy/loose match on ID
        target = next((s for s in sections if s['id'] == target_id), None)
        
        if not target:
             # Try loose match (stripped ID)
             # e.g. target_id="1." vs section_id="1"
             target = next((s for s in sections if s['id'].replace('.', '').strip() == str(target_id).replace('.', '').strip()), None)
             
        if not target:
             # Try fuzzy match on Original Text if available
             logger.warning(f"Target {target_id} not found by ID. Skipping Modify.")
             return
             
        # Update text
        # If new_content is simple string, replace.
        target['paragraph'].text = action['new_content']
        # Styling: Preserve style? Typically yes.
        # target['paragraph'].style = ... (existing)

    def _apply_insert(self, doc, sections: List[Dict], action: Dict):
        """
        action: { anchor_id: "...", position: "AFTER", new_content: "..." }
        """
        anchor_id = action.get('anchor_id')
        # Find target with fuzzy/loose match on ID
        anchor = next((s for s in sections if s['id'] == anchor_id), None)
        
        if not anchor:
             # Try loose match (stripped ID)
             anchor = next((s for s in sections if s['id'].replace('.', '').strip() == str(anchor_id).replace('.', '').strip()), None)
             
        if not anchor:
            logger.warning(f"Anchor {anchor_id} not found. Appending to end.")
            doc.add_paragraph(action['new_content'])
            return
            
        anchor_p = anchor['paragraph']
        
        # To insert AFTER anchor_p:
        # We actually need to find the *element* in XML and insert after it.
        # docx doesn't simplify this.
        # Pattern: p.insert_paragraph_before() exists.
        # We need to find the NEXT paragraph object and call insert_paragraph_before().
        
        next_p = self._get_next_paragraph(doc, anchor_p)
        
        if next_p:
            new_p = next_p.insert_paragraph_before(action['new_content'])
        else:
            # End of doc
            new_p = doc.add_paragraph(action['new_content'])
            
        # Copy style from anchor?
        new_p.style = anchor_p.style

    def _get_next_paragraph(self, doc, current_p):
        """Find the paragraph immediately following current_p in the document body"""
        # Linear search (slow but reliable for unmodified list)
        # Using XML navigation is faster: current_p._element.getnext()?
        # Let's stick to safe python-docx iteration if efficient enough, 
        # BUT doc.paragraphs is a property that reconstructs list.
        # Better: iterate once.
        
        # Optimization: XML traversal
        # current_p._element.getnext() returns an XML element.
        # We need to wrap it back to Paragraph if we use high-level API, 
        # or just use insert_paragraph_before on the *next* pure python object.
        
        try:
            # Find index in list
            # Note: This is O(N) but safer without low-level XML hacks for now.
            paras = doc.paragraphs
            idx = paras.index(current_p)
            if idx < len(paras) - 1:
                return paras[idx + 1]
        except ValueError:
            pass
        return None

    def _reindex_document(self, doc):
        """
        Scan document and enforce sequential numbering for matching sections.
        """
        counter = 1
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            
            # Use the new HEADING_RE
            match = self.HEADING_RE.match(text)
            
            if match and (match.group("article") or match.group("num")):
                # Logic: We only re-index "Article X" or "1.X" found by regex.
                # If regex caught "Chapter", we might skip re-indexing or handle separately.
                # For now: focus on "Điều" (Article).
                
                clause_type = "Điều"
                if match.group("article"):
                    # Extract "Điều " part? Or just force standard "Điều"
                    clause_type = "Điều"
                elif match.group("num"):
                     # Skip numerical bullets like "1.1" for now? 
                     # Re-indexing nested bullets is complex.
                     # Focus on fixing the main "Điều X" sequence.
                     continue
                else:
                    continue

                # Reconstruct content
                original_title = match.group("title") or ""
                new_text = f"{clause_type} {counter}: {original_title}"
                if not original_title:
                     # Attempt to preserve text after ID if regex didn't capture title well
                     # (Regex group 'title' captures everything after '[:.]?')
                     pass
                     
                p.text = new_text
                counter += 1

contract_modifier = ContractModifier()
