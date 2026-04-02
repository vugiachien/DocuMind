from sqlalchemy.orm import Session, joinedload
from fastapi import HTTPException
from typing import List, Optional
from app.db import models
from docx import Document
from app.modules.agreements.schemas import ContractCreate, RiskUpdateSuggestion, ContractShareCreate
from app.services.storage_service import storage_service
from app.services.document_service import document_service
from app.services.exceptions import TextReplacementError
from datetime import datetime
import io
import os
import re
import logging
from app.services.contract_modifier import contract_modifier
from sqlalchemy import or_, and_
from app.modules.notifications.service import notification_service
from app.services.audit_service import AuditService

logger = logging.getLogger(__name__)

class ContractService:
    def check_permission(self, db: Session, contract_id: str, user: models.User, required: str = 'view') -> models.Agreement:
        """
        Check if user has permission (view/edit) on the agreement.
        Returns agreement object if allowed, raises HTTPException if denied.
        """
        agreement = db.query(models.Agreement).options(
            joinedload(models.Agreement.shares)
        ).filter(models.Agreement.id == contract_id).first()
        
        if not agreement:
            raise HTTPException(status_code=404, detail="Agreement not found")

        # 1. Admin & Owner always have full access
        if user.role == "admin" or agreement.ownerId == user.id:
            return agreement

        # 2. Check Shares
        has_access = False
        user_permission = None

        for share in agreement.shares:
            # Check User Share
            if share.sharedType == 'user' and share.targetId == user.id:
                 user_permission = share.permission
                 has_access = True
                 break # User share takes precedence (or just take the first match)
            
            # Check Department Share
            if share.sharedType == 'department' and user.departmentId and share.targetId == user.departmentId:
                 user_permission = share.permission
                 has_access = True
                 # Don't break immediately, maybe user share has higher diff permission?
                 # ideally explicit user share > dept share.
        
        if not has_access:
             raise HTTPException(status_code=403, detail="Not authorized to access this agreement")

        # 3. Check specific permission level
        # if required='edit', need 'edit'. if required='view', 'view' or 'edit' is fine.
        if required == 'edit' and user_permission != 'edit':
             raise HTTPException(status_code=403, detail="Write permission required")
             
        return agreement

    def get_contracts(
        self, 
        db: Session, 
        user: models.User, 
        status: Optional[str] = None,
        include_deleted: bool = False
    ) -> List[models.Agreement]:
        """
        Get all agreements. 
        - Admin sees all (including deleted if include_deleted=True).
        - Regular users see pending/owned or SHARED agreements.
        - Deleted agreements are excluded by default.
        """
        # Eager load shares for permission calculation
        query = db.query(models.Agreement).options(joinedload(models.Agreement.shares))
        
        # Filter out soft-deleted agreements (unless admin explicitly requests them)
        if not include_deleted or user.role != "admin":
            query = query.filter(models.Agreement.deleted_at.is_(None))
        
        # RBAC: Filter by owner OR shared
        if user.role != "admin":
            owner_condition = models.Agreement.ownerId == user.id
            
            # Check User Share
            user_share_condition = models.Agreement.shares.any(
                and_(
                    models.ContractShare.sharedType == 'user',
                    models.ContractShare.targetId == user.id
                )
            )
            
            # Check Department Share
            dept_share_condition = False
            if user.departmentId:
                dept_share_condition = models.Agreement.shares.any(
                    and_(
                        models.ContractShare.sharedType == 'department',
                        models.ContractShare.targetId == user.departmentId
                    )
                )
            
            query = query.filter(or_(owner_condition, user_share_condition, dept_share_condition))
        
        # Filter by status
        if status:
            query = query.filter(models.Agreement.status == status)
        
        # Always show most recently updated agreements first
        query = query.order_by(models.Agreement.updatedAt.desc())
        
        agreements = query.all()
        
        # Determine Permission for each agreement
        for agreement in agreements:
            if user.role == "admin":
                agreement.currentUserPermission = "admin"
            elif agreement.ownerId == user.id:
                 agreement.currentUserPermission = "owner"
            else:
                 # Check shares (already loaded)
                 perm = "none"
                 for share in agreement.shares:
                     if share.sharedType == 'user' and share.targetId == user.id:
                         perm = share.permission # User share overrides
                         break
                     if share.sharedType == 'department' and user.departmentId and share.targetId == user.departmentId:
                         if perm == "none" or share.permission == 'edit': # Edit overrides view if multiple depts? assume yes
                             perm = share.permission
                 
                 # Map 'none' to 'view' if they are here via RBAC query but explicit share logic missed it (fallback)
                 # Actually if they passed the query filter, they have access. 
                 # If explicit logic yields 'none', it implies logic mismatch.
                 if perm == "none":
                     perm = "view" # Fallback safe default
                     
                 agreement.currentUserPermission = perm
        
        return agreements

    def get_contract(
        self, 
        db: Session, 
        contract_id: str, 
        user: models.User,
        include_deleted: bool = False
    ) -> models.Agreement:
        """
        Get specific agreement with details (Versions, Findings).
        - Enforces ownership OR share check (View Access).
        - Deleted agreements are hidden unless admin explicitly requests them.
        """
        query = db.query(models.Agreement).options(
            joinedload(models.Agreement.versions),
            joinedload(models.Agreement.findings),
            joinedload(models.Agreement.shares) 
        ).filter(models.Agreement.id == contract_id)
        
        # Filter out soft-deleted agreements (unless admin explicitly requests them)
        if not include_deleted or user.role != "admin":
            query = query.filter(models.Agreement.deleted_at.is_(None))
        
        agreement = query.first()
        
        if not agreement:
            raise HTTPException(status_code=404, detail="Agreement not found")
            
        # RBAC Check (Inline to avoid double query)
        if user.role != "admin" and agreement.ownerId != user.id:
            has_access = False
            for share in agreement.shares:
                if (share.sharedType == 'user' and share.targetId == user.id) or \
                   (share.sharedType == 'department' and user.departmentId and share.targetId == user.departmentId):
                    has_access = True
                    break
            
            if not has_access:
                raise HTTPException(status_code=403, detail="Not authorized to access this agreement")
            
        # Sort versions by uploadedAt desc
        if agreement.versions:
            agreement.versions.sort(key=lambda x: x.uploadedAt, reverse=True)

        # Sort findings by section number
        if agreement.findings:
            def get_section_num(finding):
                try:
                    if finding.section and finding.section.startswith("AUTO-"):
                        return int(finding.section.split(" ")[0].replace("AUTO-", ""))
                    return 99999
                except:
                    return 99999
            agreement.findings.sort(key=get_section_num)
            
        # Determine Permission
        if user.role == "admin":
            agreement.currentUserPermission = "admin"
        elif agreement.ownerId == user.id:
             agreement.currentUserPermission = "owner"
        else:
             # Check shares (already loaded)
             perm = "none"
             for share in agreement.shares:
                 if share.sharedType == 'user' and share.targetId == user.id:
                     perm = share.permission # User share overrides
                     break
                 if share.sharedType == 'department' and user.departmentId and share.targetId == user.departmentId:
                     if perm == "none": perm = share.permission # Use department if not set
             
             agreement.currentUserPermission = perm

        return agreement

    def share_contract(self, db: Session, contract_id: str, share_data: ContractShareCreate, user: models.User):
        """
        Share a agreement with a user or department.
        Only Owner or Admin can share.
        """
        agreement = db.query(models.Agreement).filter(models.Agreement.id == contract_id).first()
        if not agreement:
            raise HTTPException(status_code=404, detail="Agreement not found")
        
        # Strict Ownership Check for SHARING (Viewer/Editor cannot share, ONLY OWNER)
        if user.role != "admin" and agreement.ownerId != user.id:
             raise HTTPException(status_code=403, detail="Only owner can share agreement")

        # Check existing
        existing = db.query(models.ContractShare).filter(
            models.ContractShare.agreementId == agreement.id,
            models.ContractShare.sharedType == share_data.sharedType,
            models.ContractShare.targetId == share_data.targetId
        ).first()
        
        if existing:
            # Update permission if exists
            existing.permission = share_data.permission
            db.commit()
            db.refresh(existing)
            return existing

        share = models.ContractShare(
            agreementId=agreement.id,
            sharedType=share_data.sharedType,
            targetId=share_data.targetId,
            permission=share_data.permission,
            sharedBy=user.id,
            sharedAt=datetime.utcnow()
        )
        db.add(share)
        db.commit()
        db.refresh(share)
        
        # Notify Recipient
        sender_name = user.full_name or user.username
        
        # Construct payload for optimistic UI
        notification_payload = {
            "entity": "agreement",
            "action": "share",
            "data": {
                "id": agreement.id,
                "name": agreement.name,
                "partnerName": agreement.partner.name if agreement.partner else "",
                "contractTypeName": agreement.contract_type.name if agreement.contract_type else "",
                "status": agreement.status,
                "ownerId": agreement.ownerId, 
                "currentUserPermission": share.permission,
                "updatedAt": agreement.updatedAt.isoformat() if agreement.updatedAt else datetime.utcnow().isoformat(),
                "createdBy": agreement.createdBy
            }
        }

        if share.sharedType == 'user':
            notification_service.create_notification(
                db,
                share.targetId,
                "Agreement Shared",
                f"{sender_name} shared agreement '{agreement.name}' with you.",
                link=f"/agreements/{agreement.id}",
                payload=notification_payload
            )
        elif share.sharedType == 'department':
            # Notify all users in the department
            dept_users = db.query(models.User).filter(models.User.departmentId == share.targetId).all()
            for dept_user in dept_users:
                 # Don't notify the sender if they are in the department
                if dept_user.id == user.id:
                    continue
                    
                notification_service.create_notification(
                    db,
                    dept_user.id, # Valid User ID
                    "Department Agreement Share",
                    f"{sender_name} shared agreement '{agreement.name}' with your department.",
                    link=f"/agreements/{agreement.id}",
                    payload=notification_payload
                )
            
        return share

    def _publish_event(self, event_name: str, contract_id: str, **extra):
        """
        Helper to publish SSE event directly to Redis
        """
        import redis
        import json
        try:
            payload = {
                "contract_id": str(contract_id),
                "event": event_name,
                "timestamp": datetime.utcnow().isoformat()
            }
            payload.update(extra)
            
            # Connect to Redis
            r = redis.from_url(os.getenv("REDIS_URL", "redis://localhost:6390/0"))
            r.publish("contract_updates", json.dumps(payload))
        except Exception as e:
            logger.error(f"Failed to publish event {event_name}: {e}")

    def revoke_share(self, db: Session, contract_id: str, share_id: str, user: models.User):
        """
        Revoke a share.
        """
        agreement = db.query(models.Agreement).filter(models.Agreement.id == contract_id).first()
        if not agreement:
            raise HTTPException(status_code=404, detail="Agreement not found")
        
        if user.role != "admin" and agreement.ownerId != user.id:
             raise HTTPException(status_code=403, detail="Only owner can manage shares")
              
        share = db.query(models.ContractShare).filter(
            models.ContractShare.id == share_id,
            models.ContractShare.agreementId == contract_id
        ).first()
        
        if not share:
            raise HTTPException(status_code=404, detail="Share not found")
            
        # Capture details before deletion for notification
        target_id = share.targetId
        shared_type = share.sharedType
        contract_name = agreement.name
        revoker_name = user.full_name or user.username
        
        db.delete(share)
        db.commit()
        
        # --- Notification & SSE Logic ---
        try:
            notification_payload = {
                "entity": "agreement",
                "action": "revoke",
                "data": {
                    "id": agreement.id,
                    "name": agreement.name
                }
            }
            
            if shared_type == 'user':
                # invalidating cache or notifying user
                 notification_service.create_notification(
                    db,
                    target_id,
                    "Access Revoked",
                    f"{revoker_name} has revoked your access to agreement '{contract_name}'.",
                    type="warning",
                    link=None, # No link since access lost
                    payload=notification_payload
                )
                 # Publish specific event for this user to remove from list
                 self._publish_event("contract_revoked", agreement.id, target_user_id=target_id)
                 
            elif shared_type == 'department':
                # Notify all dept users
                dept_users = db.query(models.User).filter(models.User.departmentId == target_id).all()
                for dept_user in dept_users:
                    if dept_user.id == user.id: continue
                    
                    notification_service.create_notification(
                        db,
                        dept_user.id,
                        "Access Revoked",
                        f"{revoker_name} revoked department access to agreement '{contract_name}'.",
                        type="warning",
                        link=None,
                        payload=notification_payload
                    )
                # Publish event for department
                self._publish_event("contract_revoked", agreement.id, target_dept_id=target_id)
                
        except Exception as e:
            logger.error(f"Failed to send revoke notifications: {e}")
            # Don't fail the request since revocation succeeded
            
        return True

    def delete_contract(self, db: Session, contract_id: str, user: models.User, hard_delete: bool = False):
        """
        Delete agreement (soft delete by default).
        ONLY OWNER/ADMIN can delete. Even 'edit' permission cannot delete.
        
        Args:
            hard_delete: If True, permanently delete. If False, soft delete (default).
        """
        agreement = db.query(models.Agreement).filter(models.Agreement.id == contract_id).first()
        if not agreement:
            raise HTTPException(status_code=404, detail="Agreement not found")

        # Check if already soft-deleted
        if agreement.is_deleted and not hard_delete:
            raise HTTPException(status_code=400, detail="Agreement is already deleted")

        # Strict Ownership Check
        if user.role != "admin" and agreement.ownerId != user.id:
            raise HTTPException(status_code=403, detail="Only owner or admin can delete agreement")
        
        if hard_delete:
            # Hard delete - cleanup ALL files from MinIO (main file + all version files)
            files_to_delete = set()
            if agreement.fileUrl:
                files_to_delete.add(agreement.fileUrl)
            
            # Collect version file URLs
            versions = db.query(models.ContractVersion).filter(
                models.ContractVersion.agreementId == agreement.id
            ).all()
            for ver in versions:
                if ver.fileUrl:
                    files_to_delete.add(ver.fileUrl)
            
            for file_url in files_to_delete:
                try:
                    storage_service.delete_file(file_url)
                except Exception as e:
                    logger.warning(f"Failed to delete file from MinIO ({file_url}): {e}")
            
            db.delete(agreement)
        else:
            # Soft delete - mark as deleted
            agreement.soft_delete(deleted_by_user_id=user.id)
        
        db.commit()
        return True
    
    def restore_contract(self, db: Session, contract_id: str, user: models.User):
        """
        Restore a soft-deleted agreement.
        ONLY ADMIN can restore.
        """
        agreement = db.query(models.Agreement).filter(models.Agreement.id == contract_id).first()
        if not agreement:
            raise HTTPException(status_code=404, detail="Agreement not found")
        
        if not agreement.is_deleted:
            raise HTTPException(status_code=400, detail="Agreement is not deleted")
        
        # Only admin can restore
        if user.role != "admin":
            raise HTTPException(status_code=403, detail="Only admin can restore deleted agreements")
        
        agreement.restore()
        db.commit()
        return agreement
    def create_contract(self, db: Session, contract_in: ContractCreate, user: models.User) -> models.Agreement:
        """
        Create a new agreement.
        - Validates Partner and Agreement Type existence.
        - Automatically assigns owner.
        """
        # Validate Partner
        partner = db.query(models.Partner).filter(models.Partner.id == contract_in.partnerId).first()
        if not partner:
            raise HTTPException(status_code=404, detail="Partner not found")
        
        # Validate Agreement Type
        contract_type = db.query(models.ContractType).filter(models.ContractType.id == contract_in.agreementTypeId).first()
        if not contract_type:
            raise HTTPException(status_code=404, detail="Agreement type not found")
        
        # Validate AuditPolicy (Auto-select if missing)
        if contract_in.auditPolicyId:
            audit_policy = db.query(models.AuditPolicy).filter(models.AuditPolicy.id == contract_in.auditPolicyId).first()
            if not audit_policy:
                raise HTTPException(status_code=404, detail="AuditPolicy not found")
        else:
            # Auto-select the latest active audit_policy for this agreement type
            suggested_playbook = db.query(models.AuditPolicy).filter(
                models.AuditPolicy.agreementTypeId == contract_in.agreementTypeId,
                models.AuditPolicy.status == 'active'
            ).order_by(models.AuditPolicy.uploadedAt.desc()).first()
            
            if suggested_playbook:
                contract_in.auditPolicyId = suggested_playbook.id
                logger.info(f"🤖 Auto-selected AuditPolicy '{suggested_playbook.name}' for Agreement Type '{contract_type.name}'")
        
        # Generate unique agreement number
        import uuid
        if not contract_in.contractNumber:
            contract_number = f"CNT-{uuid.uuid4().hex[:12].upper()}"
        else:
            contract_number = contract_in.contractNumber

        new_contract = models.Agreement(
            contractNumber=contract_number,
            name=contract_in.name,
            partnerId=contract_in.partnerId,
            agreementTypeId=contract_in.agreementTypeId,
            auditPolicyId=contract_in.auditPolicyId,
            ownerId=user.id,
            status='draft',
            value=contract_in.value,
            effectiveDate=contract_in.effectiveDate,
            expiryDate=contract_in.expiryDate,
            createdBy=user.username,
            fileUrl=None
        )
        
        try:
            db.add(new_contract)
            db.commit()
            db.refresh(new_contract)
        except Exception as e:
            db.rollback()
            if "duplicate key" in str(e).lower():
                raise HTTPException(status_code=400, detail=f"Agreement number '{contract_number}' already exists.")
            raise HTTPException(status_code=500, detail=str(e))
        
        return new_contract

    def update_contract(self, db: Session, contract_id: str, contract_in: ContractCreate, user: models.User) -> models.Agreement:
        """
        Update agreement details.
        """
        agreement = self.check_permission(db, contract_id, user, required='edit')
        
        # Update fields
        agreement.name = contract_in.name
        agreement.partnerId = contract_in.partnerId
        agreement.agreementTypeId = contract_in.agreementTypeId
        agreement.auditPolicyId = contract_in.auditPolicyId
        if contract_in.contractNumber:
            agreement.contractNumber = contract_in.contractNumber
        agreement.value = contract_in.value
        agreement.effectiveDate = contract_in.effectiveDate
        agreement.expiryDate = contract_in.expiryDate
        
        try:
            db.commit()
            db.refresh(agreement)
        except Exception as e:
            db.rollback()
            raise HTTPException(status_code=500, detail=str(e))
            
        return agreement

    def process_new_contract_file(
        self,
        db: Session,
        agreement: models.Agreement,
        file_path: str,
        user_id: str,
    ) -> dict:
        """
        Process a newly uploaded/converted DOCX file.
        1. Download file content from MinIO
        2. Check for template URL in ContractType
        3. If template exists, calculate cosine similarity
        4. Set isTemplateBased and templateSimilarity
        5. Create ContractVersion records (v0.0 and v0.1 for template-based, v0.0 otherwise)
        """
        from app.services.template_matcher import compute_similarity, is_template_based
        import uuid

        # Check for template
        template_url = None
        if agreement.agreementTypeId:
            contract_type = db.query(models.ContractType).filter(
                models.ContractType.id == agreement.agreementTypeId
            ).first()
            if contract_type:
                template_url = contract_type.templateUrl

        similarity_score = 0.0
        if template_url:
            try:
                file_content = storage_service.download_file(file_path)
                template_content = storage_service.download_file(template_url)
                
                upload_text = document_service.extract_text(file_content)
                template_text = document_service.extract_text(template_content)
                
                similarity_score = compute_similarity(template_text, upload_text)
                logger.info(f"📊 Template similarity for agreement {agreement.id}: {similarity_score:.3f}")
            except Exception as e:
                logger.warning(f"Similarity check failed (non-blocking): {e}")

        is_based = is_template_based(similarity_score) if template_url else False
        agreement.isTemplateBased = is_based
        agreement.templateSimilarity = round(similarity_score, 4) if template_url else None

        # Clear existing versions if any
        db.query(models.ContractVersion).filter(
            models.ContractVersion.agreementId == agreement.id
        ).delete()

        if is_based and template_url:
            # v0.0 = template baseline
            v0 = models.ContractVersion(
                id=str(uuid.uuid4()),
                agreementId=agreement.id,
                version="v0.0",
                fileUrl=template_url,
                createdBy=user_id,
                changes="Template baseline (auto-created)",
                versionType="template",
            )
            db.add(v0)
            
            # v0.1 = uploaded file
            v1 = models.ContractVersion(
                id=str(uuid.uuid4()),
                agreementId=agreement.id,
                version="v0.1",
                fileUrl=file_path,
                createdBy=user_id,
                changes="Initial upload",
                versionType="upload",
            )
            db.add(v1)
            agreement.fileUrl = file_path
            agreement.currentVersion = "v0.1"
            
            msg = f"Template match detected ({similarity_score:.0%}). Versions v0.0 + v0.1 created."
        else:
            # v0.0 = uploaded file (single baseline)
            v0 = models.ContractVersion(
                id=str(uuid.uuid4()),
                agreementId=agreement.id,
                version="v0.0",
                fileUrl=file_path,
                createdBy=user_id,
                changes="Initial upload",
                versionType="upload",
            )
            db.add(v0)
            agreement.fileUrl = file_path
            agreement.currentVersion = "v0.0"
            
            msg = "No template match. Version v0.0 created (uploaded file)."

        agreement.status = "draft"
        db.commit()
        db.refresh(agreement)
        
        return {
            "isTemplateBased": agreement.isTemplateBased,
            "templateSimilarity": agreement.templateSimilarity,
            "currentVersion": agreement.currentVersion,
            "message": msg
        }

    def analyze_contract(self, db: Session, contract_id: str, user: models.User, background_tasks, full_context_mode: bool = False) -> models.Agreement:
        """
        Trigger AI analysis for a agreement.
        """
        agreement = self.check_permission(db, contract_id, user, required='view') # Viewers can analyze
        
        if agreement.status == "converting":
            raise HTTPException(status_code=400, detail="Agreement is still converting. Please wait.")

        if not agreement.fileUrl:
            raise HTTPException(status_code=400, detail="No file uploaded for this agreement")
        
        # Create initial version record if not exists
        from datetime import datetime
        current_ver = db.query(models.ContractVersion).filter(
            models.ContractVersion.agreementId == agreement.id, 
            models.ContractVersion.version == agreement.currentVersion
        ).first()
        
        if not current_ver:
            current_ver = models.ContractVersion(
                agreementId=agreement.id,
                version=agreement.currentVersion,
                fileUrl=agreement.fileUrl,
                uploadedAt=agreement.updatedAt or datetime.utcnow(),
                createdBy=agreement.createdBy,
                changes="Initial Upload"
            )
            db.add(current_ver)
            db.commit()
            db.refresh(current_ver)
            
        # Set status to processing
        current_ver.processingStatus = "pending"
        agreement.status = "processing"
        db.commit()
        
        # Enqueue Task – pass full_context_mode flag to worker
        from app.worker import process_upload_task
        background_tasks.add_task(process_upload_task, current_ver.id, full_context_mode)
        
        logger.info(f"✅ Enqueued analysis task in BackgroundTasks (full_context_mode={full_context_mode})")
        
        return agreement


    
    def update_risk_suggestion(self, db: Session, contract_id: str, risk_id: str, payload: RiskUpdateSuggestion, user: models.User):
        """
        Update the suggested_text for a finding before applying.
        """
        agreement = self.check_permission(db, contract_id, user, required='edit')
        
        finding = db.query(models.Finding).filter(
            models.Finding.agreementId == contract_id,
            models.Finding.id == risk_id
        ).first()
        
        if not finding:
            raise HTTPException(status_code=404, detail="Finding not found")
            
        finding.suggested_text = payload.updated_text
        finding.auto_fixable = True  # ✅ Manual edit implies user approval/fix
        db.commit()
        db.refresh(finding)
        return finding

    def accept_risk_recommendation(self, db: Session, contract_id: str, risk_id: str, current_version: str, user: models.User):
        """
        Accept single finding recommendation. Wraps batch logic.
        """
        return self.accept_risk_batch(db, contract_id, [risk_id], current_version, user)

    def accept_risk_batch(self, db: Session, contract_id: str, risk_ids: List[str], current_version: str, user: models.User):
        """
        Accept multiple finding recommendations at once.
        Applies fixes to DOCX, updates MinIO, updates DB versions.
        """
        # 1. Fetch Data & Lock
        # Using check_permission here might break validation if we need lock first.
        # But RBAC check should happen early.
        # Let's check permission first without lock, then query with lock? 
        # Or just do RBAC check manually here since logic is complex with locking.
        
        # Using check_permission (no lock) primarily for authorization
        self.check_permission(db, contract_id, user, required='edit')

        agreement = db.query(models.Agreement).filter(models.Agreement.id == contract_id).with_for_update().first()
        if not agreement:
            raise HTTPException(status_code=404, detail="Agreement not found")
            
        # RBAC Check is done above via check_permission

        # Check Version Conflict
        if agreement.currentVersion != current_version:
            raise HTTPException(
                status_code=409, 
                detail=f"Conflict: Agreement has been updated to {agreement.currentVersion} by another user. Please refresh."
            )
            
        findings = db.query(models.Finding).filter(
            models.Finding.agreementId == contract_id,
            models.Finding.id.in_(risk_ids)
        ).all()
        
        if not findings:
            raise HTTPException(status_code=404, detail="No valid findings found")
            
        if not agreement.fileUrl:
            raise HTTPException(status_code=400, detail="No file associated with this agreement")
            
        if not storage_service:
            raise HTTPException(status_code=503, detail="Storage service unavailable")

        try:
            # Determine New Version
            current_ver_str = agreement.currentVersion
            try:
                ver_clean = current_ver_str.lower().lstrip('v')
                major, minor = map(int, ver_clean.split('.'))
                new_minor = minor + 1
                new_ver_str = f"v{major}.{new_minor}"
            except Exception as e:
                logger.error(f"Error parsing version {current_ver_str}: {e}")
                # Fallback: strip any non-numeric prefix, default to v1.1
                import re as _re
                nums = _re.findall(r'\d+', current_ver_str)
                if nums:
                    new_ver_str = f"v{nums[0]}.{int(nums[-1]) + 1}"
                else:
                    new_ver_str = "v1.1"
                
            # History Snapshot (if needed)
            current_version_record = db.query(models.ContractVersion).filter(
                models.ContractVersion.agreementId == agreement.id,
                models.ContractVersion.version == current_ver_str
            ).first()
            
            if not current_version_record:
                history_record = models.ContractVersion(
                    agreementId=agreement.id,
                    version=current_ver_str,
                    fileUrl=agreement.fileUrl,
                    uploadedAt=agreement.updatedAt or datetime.utcnow(),
                    createdBy=agreement.createdBy,
                    changes="Original version"
                )
                db.add(history_record)

            # Download File
            file_content = storage_service.download_file(agreement.fileUrl)
            
            # SIMPLE TEXT REPLACEMENT APPROACH (more reliable than ID-based)
            # Filter: Only process "modification" findings (skip "recommendation" findings)
            # SIMPLE TEXT REPLACEMENT APPROACH (more reliable than ID-based)
            # Filter: Only process "modification" findings (skip "recommendation" findings)
            replacements = []
            logs = []
            skipped_recommendations = 0
            skipped_missing = 0
            processed_modifications = 0
            
            for finding in findings:
                # Skip recommendation-only findings (require manual edit)
                if getattr(finding, 'risk_type', 'modification') == 'recommendation':
                    skipped_recommendations += 1
                    logs.append(f"Skipped recommendation finding: {finding.id}")
                    continue
                    
                # Process modification findings
                if finding.original_text and finding.suggested_text and finding.original_text != "(Missing Clause)":
                    replacements.append((finding.original_text, finding.suggested_text))
                    logs.append(f"Replacing: {finding.original_text[:30]}...")
                    processed_modifications += 1
                elif not finding.original_text or finding.original_text == "(Missing Clause)":
                     # Skip missing clauses (Require manual edit per user request)
                     skipped_missing += 1
                     logs.append(f"Skipped missing clause finding: {finding.id}")
                else:
                    logs.append(f"Skipping finding {finding.id}: missing suggested_text")

            if not replacements:
                # Graceful handling: Return success with message instead of Error 400
                msg = "No changes applied found."
                if skipped_missing > 0:
                     msg = "Selected items require manual editing (Missing Clauses)."
                elif skipped_recommendations > 0:
                    msg = "Selected items require manual editing (Recommendations)."
                
                return {
                    "message": msg, 
                    "newVersion": agreement.currentVersion, 
                    "processed": 0,
                    "skipped_recommendations": skipped_recommendations + skipped_missing
                }

            # Execute Text Replacement
            try:
                # Use document_service's batch replacement for efficiency
                modified_content, failed_targets = document_service.replace_multiple_texts_in_docx(file_content, replacements)
                
                success_count = len(replacements) - len(failed_targets)
                logs.append(f"Applied {success_count}/{len(replacements)} text replacements successfully")
                
                if failed_targets:
                     logger.warning(f"Failed to replace {len(failed_targets)} items")
            
                # Update Descriptions (only for successfully processed modifications)
                for finding in findings:
                    # Skip if finding is recommendation 
                    if getattr(finding, 'risk_type', 'modification') == 'recommendation':
                        continue
                        
                    # Check if this finding failed
                    if finding.original_text in failed_targets:
                        # Maybe mark as failed? Or just don't mark resolved.
                        # Do NOT mark as resolved.
                        continue
                     
                    if not finding.original_text or finding.original_text == "(Missing Clause)":
                        continue

                    if "[RESOLVED]" not in finding.description:
                         finding.description = f"[RESOLVED] {finding.description}"
                         
            except Exception as e:
                logger.error(f"Text replacement failed: {e}")
                raise HTTPException(status_code=500, detail=f"Document modification failed: {str(e)}")

            # Generate New Filename
            path_parts = agreement.fileUrl.split('/')
            original_filename = path_parts[-1]
            directory = "/".join(path_parts[:-1])
            name, ext = os.path.splitext(original_filename)
            base_name = re.sub(r'_v\d+(\.\d+)?$', '', name)
            new_filename = f"{base_name}_{new_ver_str}{ext}"
            new_file_path = f"{directory}/{new_filename}"
            
            # Upload
            file_obj = io.BytesIO(modified_content)
            storage_service.client.put_object(
                storage_service.bucket,
                new_file_path,
                file_obj,
                len(modified_content),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Update DB (New Version)
            try:
                extracted_text = document_service.extract_text(modified_content)
                html_preview = document_service.convert_docx_to_html(modified_content) # ✅ Generate HTML Preview immediately
            except Exception as e:
                logger.warning(f"Failed to generate preview for auto-apply: {e}")
                extracted_text = None
                html_preview = None

            new_version_record = models.ContractVersion(
                agreementId=agreement.id,
                version=new_ver_str,
                fileUrl=new_file_path,
                uploadedAt=datetime.utcnow(),
                createdBy=user.username,
                changes=f"Auto-applied fixes: {success_count} text replacements.",
                extractedText=extracted_text,
                htmlPreview=html_preview  # ✅ Save HTML Preview
            )
            db.add(new_version_record)
            
            agreement.currentVersion = new_ver_str
            agreement.fileUrl = new_file_path
            agreement.updatedAt = datetime.utcnow()
            agreement.status = 'update'
            # --- AUDIT LOGGING ---
            AuditService.log_activity(
                db,
                user.id,
                "RISK_BATCH_APPLY",
                "AGREEMENT",
                agreement.id,
                details={
                    "newVersion": new_ver_str,
                    "processed": success_count,
                    "skipped": skipped_recommendations,
                    "logs": logs
                }
            )

            db.commit()

            failed_count = len(failed_targets) if failed_targets else 0
            return {
                "message": "Findings accepted and file updated", 
                "newVersion": new_ver_str, 
                "processed": success_count,
                "skipped_recommendations": skipped_recommendations,
                "skipped_missing": skipped_missing,
                "failed_count": failed_count
            }

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error in accept_risk_batch: {e}")
            db.rollback()
            raise HTTPException(status_code=500, detail=str(e))

    def create_manual_version(self, db: Session, contract_id: str, content_text: str, changes_desc: str, user: models.User, resolved_risk_ids: List[str] = []):
        """
        Create a new version from manually edited text content.
        Generates a new DOCX file from the text.
        """
        agreement = self.check_permission(db, contract_id, user, required='edit')
        if not agreement:
            raise HTTPException(status_code=404, detail="Agreement not found")
            
        try:
            # 1. Determine New Version Number
            current_ver_str = agreement.currentVersion
            v_parts = current_ver_str.replace('v', '').split('.')
            if len(v_parts) == 2:
                new_ver_str = f"v{v_parts[0]}.{int(v_parts[1]) + 1}"
            else:
                new_ver_str = f"{current_ver_str}.1"
                
            # 2. Generate DOCX from Text
            doc = Document()
            
            # Check if content is HTML (from ReactQuill)
            is_html = False
            
            # --- SECURITY FIX: Sanitize Input (XSS Prevention) ---
            if isinstance(content_text, str) and ("<p" in content_text or "<br" in content_text):
                try:
                    import bleach
                    # Allow tags used by TinyMCE/RichText
                    allowed_tags = [
                        'p', 'br', 'strong', 'b', 'em', 'i', 'u', 'ul', 'ol', 'li', 
                        'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'span', 'div', 'a'
                    ]
                    # Allow attributes
                    allowed_attrs = {
                        '*': ['style', 'class'],
                        'a': ['href', 'target'],
                        'span': ['data-comment-id']
                    }
                    # Allow styles (essential for text alignment, colors)
                    allowed_styles = [
                        'text-align', 'color', 'background-color', 'font-size', 'font-weight', 
                        'text-decoration', 'margin', 'padding', 'padding-left', 'list-style-type'
                    ]
                    
                    content_text = bleach.clean(
                        content_text, 
                        tags=allowed_tags, 
                        attributes=allowed_attrs, 
                        styles=allowed_styles, 
                        strip=True
                    )
                except ImportError:
                    logger.warning("Bleach not installed. Skipping sanitization (Security Finding)")
                except Exception as e:
                    logger.error(f"Sanitization failed: {e}")
            # -----------------------------------------------------

            final_extracted_text = content_text

            if isinstance(content_text, str) and ("<p" in content_text or "<br" in content_text):
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(content_text, 'html.parser')
                    is_html = True
                    final_extracted_text = soup.get_text(separator='\n')
                    
                    # Iterate elements roughly in order
                    # Note: find_all finds nested tags too, so we should iterate children of body if possible, or just top level
                    # But Quill wraps everything.
                    # Simple heuristic: find all block elements
                    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'ul', 'ol']):
                        if element.name == 'p':
                            p = doc.add_paragraph()
                            if element.get_text(strip=True) == "":
                                continue # Skip empty
                            # Iterate children for formatting
                            for child in element.children:
                                if child.name == None: # NavigableString
                                    p.add_run(str(child))
                                elif child.name in ['b', 'strong']:
                                    run = p.add_run(child.get_text())
                                    run.bold = True
                                elif child.name in ['i', 'em']:
                                    run = p.add_run(child.get_text())
                                    run.italic = True
                                elif child.name == 'u':
                                    run = p.add_run(child.get_text())
                                    run.underline = True
                                elif child.name == 'br':
                                    p.add_run('\n')
                                else:
                                    p.add_run(child.get_text())
                                    
                        elif element.name.startswith('h'):
                             level = int(element.name[1])
                             doc.add_heading(element.get_text(), level=level)
                        elif element.name in ['ul', 'ol']:
                             for li in element.find_all('li'):
                                 # 'List Bullet' or 'List Number' usually available in default template
                                 # Using default param logic
                                 try:
                                     style = 'List Bullet' if element.name == 'ul' else 'List Number'
                                     doc.add_paragraph(li.get_text(), style=style)
                                 except:
                                     doc.add_paragraph(li.get_text()) # Fallback

                except Exception as e:
                    logger.warning(f"HTML parsing failed: {e}")
                    is_html = False

            if not is_html:
                # Split by lines to preserve basic paragraphs
                for line in content_text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
                    elif not line:
                        pass
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            file_content = output.read()
            
            # 3. Upload File
            # Generate New Filename
            if agreement.fileUrl:
                path_parts = agreement.fileUrl.split('/')
                original_filename = path_parts[-1]
                directory = "/".join(path_parts[:-1])
                name, ext = os.path.splitext(original_filename)
                base_name = re.sub(r'_v\d+(\.\d+)?$', '', name)
                new_filename = f"{base_name}_{new_ver_str}{ext}"
                new_file_path = f"{directory}/{new_filename}"
            else:
                # Default path if no previous file
                new_filename = f"{agreement.contractNumber}_{new_ver_str}.docx"
                new_file_path = f"agreements/{agreement.partnerId or 'unknown'}/{new_filename}"

            storage_service.client.put_object(
                storage_service.bucket,
                new_file_path,
                io.BytesIO(file_content),
                len(file_content),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # 4. Save Version Record
            # Extract plain text for search/analysis, but keep HTML for preview
            plain_text_version = final_extracted_text
            if is_html:
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(content_text, 'html.parser')
                    # strip=True removes surrounding whitespace from text nodes
                    plain_text_version = soup.get_text(separator='\n', strip=True) 
                    
                    # Normalize excessive newlines (max 2)
                    plain_text_version = re.sub(r'\n{3,}', '\n\n', plain_text_version)
                except:
                    plain_text_version = content_text

            new_version_record = models.ContractVersion(
                agreementId=agreement.id,
                version=new_ver_str,
                fileUrl=new_file_path,
                uploadedAt=datetime.utcnow(),
                createdBy=user.username,
                changes=changes_desc,
                extractedText=plain_text_version,
                htmlPreview=content_text if is_html else None # ✅ Save HTML for perfect preview
            )
            db.add(new_version_record)
            
            # 5. Update Agreement
            agreement.currentVersion = new_ver_str
            agreement.fileUrl = new_file_path
            agreement.updatedAt = datetime.utcnow()
            agreement.status = 'update'
            
            # Update Resolved Findings
            if resolved_risk_ids:
                risks_to_resolve = db.query(models.Finding).filter(
                    models.Finding.agreementId == contract_id,
                    models.Finding.id.in_(resolved_risk_ids)
                ).all()
                for finding in risks_to_resolve:
                    if "[RESOLVED]" not in finding.description:
                        finding.description = f"[RESOLVED] {finding.description}"
            
            db.commit()
            db.refresh(agreement)
            
            return new_version_record
            
        except Exception as e:
            logger.error(f"Error creating manual version: {e}")
            db.rollback()
            raise HTTPException(status_code=500, detail=str(e))

contract_service = ContractService()
